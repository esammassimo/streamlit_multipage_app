import re
import html
from pathlib import Path
from typing import Dict, List, Tuple, Any, Optional, Iterable, Union

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

try:
    import openpyxl
    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False


# =========================
# CONFIG
# =========================

OUTPUT_META_LABELS = ["Title", "Description", "URL", "Territories", "Target Keyword"]

# Varianti input -> chiave output (case-insensitive)
INPUT_KEY_MAP = {
    "title": "Title",
    "seo title": "Title",
    "meta title": "Title",

    "description": "Description",
    "meta description": "Description",

    "url": "URL",

    "territories": "Territories",
    "territory": "Territories",

    "target keyword": "Target Keyword",
    "target keywords": "Target Keyword",
    "kw": "Target Keyword",
    "keyword": "Target Keyword",
    "primary keyword": "Target Keyword",

    "h1": "H1",
}

# riconosci “Testo:” con varianti
TESTO_KEYS = {"testo", "content", "article", "body", "testo articolo"}


# =========================
# EXCEL LOOKUP (immagini + carosello)
# =========================

# Nome del foglio e indici colonne (0-based)
PED_SHEET_NAME = "PED 2026 - Short View"
PED_COL_SLUG   = 7   # H  – URL/SLUG
PED_COL_H1     = 8   # I  – H1
PED_COL_NIMGS  = 16  # Q  – N° image for articles
PED_COL_PRODS  = 17  # R  – ID Products for Carousel


def _parse_n_images(raw) -> int:
    """
    Estrae il numero di immagini da valori come:
      4.0  ->  4
      '7 (sceglierne 5)'  ->  5   (prende il primo numero tra parentesi)
      '4'  ->  4
    """
    if raw is None:
        return 0
    s = str(raw).strip()
    # cerca prima numero tra parentesi, es. "(sceglierne 5)"
    m = re.search(r"\(.*?(\d+).*?\)", s)
    if m:
        return int(m.group(1))
    # altrimenti prende il primo numero nella stringa
    m = re.search(r"(\d+)", s)
    if m:
        return int(m.group(1))
    return 0


def _parse_product_ids(raw) -> List[str]:
    """
    Estrae lista di product ID da stringhe come:
      '01330-LAC ; 00769-LAC'  ->  ['01330-LAC', '00769-LAC']
      None                     ->  []
    """
    if not raw:
        return []
    parts = re.split(r"\s*;\s*", str(raw).strip())
    return [p.strip() for p in parts if p.strip()]


def load_ped_lookup(xlsx_path: Path) -> Dict[str, Dict[str, Any]]:
    """
    Carica il foglio PED e restituisce un dizionario:
      { slug_normalizzato: { "n_images": int, "product_ids": [str,...] }, ... }
    La chiave è lo slug (col H) normalizzato (strip + lower).
    Costruisce anche un indice secondario per H1.
    """
    if not _OPENPYXL_AVAILABLE:
        print("WARNING: openpyxl non disponibile – placeholder immagini/carosello disabilitati.")
        return {}

    if not xlsx_path or not Path(xlsx_path).exists():
        print(f"WARNING: file PED non trovato ({xlsx_path}) – placeholder disabilitati.")
        return {}

    wb = openpyxl.load_workbook(str(xlsx_path), read_only=True, data_only=True)
    if PED_SHEET_NAME not in wb.sheetnames:
        print(f"WARNING: foglio '{PED_SHEET_NAME}' non trovato nel file Excel.")
        return {}

    ws = wb[PED_SHEET_NAME]
    lookup: Dict[str, Dict[str, Any]] = {}

    for row in ws.iter_rows(min_row=2, values_only=True):
        if row is None or all(v is None for v in row):
            continue
        slug = row[PED_COL_SLUG]
        h1   = row[PED_COL_H1]
        raw_imgs  = row[PED_COL_NIMGS] if len(row) > PED_COL_NIMGS else None
        raw_prods = row[PED_COL_PRODS] if len(row) > PED_COL_PRODS else None

        entry = {
            "n_images":    _parse_n_images(raw_imgs),
            "product_ids": _parse_product_ids(raw_prods),
        }

        if slug:
            lookup[str(slug).strip().lower()] = entry
        if h1:
            lookup["__h1__" + str(h1).strip().lower()] = entry

    return lookup


def lookup_article_extras(
    lookup: Dict[str, Dict[str, Any]],
    slug: str = "",
    h1:   str = "",
) -> Dict[str, Any]:
    """
    Cerca l'entry nel dizionario PED per slug o H1.
    Ritorna { "n_images": int, "product_ids": [...] } o valori vuoti.
    """
    default = {"n_images": 0, "product_ids": []}
    if not lookup:
        return default

    if slug:
        key = slug.strip().lower()
        # rimuovi eventuale "/" iniziale o "-" iniziale (docx a volte ha "- benefici-...")
        key = key.lstrip("-/ ")
        if key in lookup:
            return lookup[key]

    if h1:
        h1_key = "__h1__" + h1.strip().lower()
        if h1_key in lookup:
            return lookup[h1_key]

    return default


# =========================
# HTML helpers
# =========================

def html_entities(s: str) -> str:
    """
    Converte caratteri speciali in HTML entities.
    - Escapa & (e altri caratteri speciali)
    - Converte virgolette tipografiche e simboli frequenti
    """
    if not s:
        return ""

    s = html.escape(s, quote=False)  # gestisce & < > e caratteri non-ascii come entities

    replacements = {
        "’": "&rsquo;",
        "‘": "&lsquo;",
        "“": "&ldquo;",
        "”": "&rdquo;",
        "–": "&ndash;",
        "—": "&mdash;",
        "…": "&hellip;",
        "à": "&agrave;",
        "è": "&egrave;",
        "é": "&eacute;",
        "ì": "&igrave;",
        "ò": "&ograve;",
        "ù": "&ugrave;",
        "À": "&Agrave;",
        "È": "&Egrave;",
        "É": "&Eacute;",
        "Ì": "&Igrave;",
        "Ò": "&Ograve;",
        "Ù": "&Ugrave;",
        "ô": "&ocirc;",
        "Ô": "&Ocirc;"
    }
    for k, v in replacements.items():
        s = s.replace(k, v)
    return s


# =========================
# DOCX helpers
# =========================

def shade_cell(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)

def set_cell_text(cell, text: str, bold: bool = False, color: Optional[RGBColor] = None, size_pt: int = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text or "")
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color is not None:
        run.font.color.rgb = color

def _normalize_key(k: str) -> str:
    k = k.strip().lower()
    k = re.sub(r"\s+", " ", k)
    return k

def _is_document_like(obj) -> bool:
    return hasattr(obj, "element") and hasattr(obj.element, "body")

def _is_cell_like(obj) -> bool:
    return hasattr(obj, "_tc")

def iter_block_items(parent) -> Iterable[Union[Paragraph, Table]]:
    """
    Itera paragraph e tabelle in ordine di apparizione.
    Supporta Document e celle.
    """
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl

    if _is_document_like(parent):
        parent_elm = parent.element.body
    elif _is_cell_like(parent):
        parent_elm = parent._tc
    else:
        raise TypeError("Parent non supportato: {}".format(type(parent)))

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# =========================
# List detection (bullets/numbered)
# =========================

def is_list_paragraph(paragraph: Paragraph) -> bool:
    """
    Ritorna True se il paragrafo è un elemento di lista (bullet o numbered).
    In Word entrambe le tipologie hanno w:numPr.
    """
    p = paragraph._p
    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        return False
    numPr = pPr.find(qn("w:numPr"))
    return numPr is not None


# =========================
# Hyperlink extraction (standard + field codes)
# =========================

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

def _runs_text(run_elm) -> str:
    """Estrae testo (concatenato) da un <w:r> (run XML)."""
    out = []
    for node in run_elm.iter():
        if node.tag.endswith("}t") and node.text:
            out.append(node.text)
    return "".join(out)

def _extract_url_from_instr(instr: str) -> str:
    """
    Estrae URL da istruzioni tipo:
      HYPERLINK "https://example.com"
      HYPERLINK https://example.com
    """
    if not instr:
        return "#"
    instr = instr.strip()

    m = re.search(r'HYPERLINK\s+"([^"]+)"', instr, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()

    m = re.search(r"HYPERLINK\s+(\S+)", instr, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip()

    return "#"

def paragraph_to_html(paragraph: Paragraph) -> str:
    """
    Converte un Paragraph in HTML-safe,
    convertendo TUTTI gli hyperlink Word in <a href="...">...</a>:
      - <w:hyperlink r:id="...">
      - <w:fldSimple w:instr="HYPERLINK ...">
      - field codes complessi: fldChar begin/separate/end + instrText
    """
    parts: List[str] = []
    p_elm = paragraph._p

    # Stato per campi complessi
    in_field = False
    field_instr = ""
    field_url = None
    after_separate = False
    field_display_parts: List[str] = []

    def flush_field():
        nonlocal in_field, field_instr, field_url, after_separate, field_display_parts
        if field_url and field_display_parts:
            visible = "".join(field_display_parts).strip()
            if visible:
                parts.append(
                    '<a href="{href}">{text}</a>'.format(
                        href=html.escape(field_url, quote=True),
                        text=html_entities(visible)
                    )
                )
        in_field = False
        field_instr = ""
        field_url = None
        after_separate = False
        field_display_parts = []

    for child in p_elm.iterchildren():
        tag = child.tag.split("}")[-1]

        # (A) Hyperlink standard
        if tag == "hyperlink":
            r_id = child.get(qn("r:id"))
            url = "#"
            if r_id:
                rel = paragraph.part.rels.get(r_id)
                if rel and getattr(rel, "target_ref", None):
                    url = rel.target_ref

            link_text_parts = []
            for grand in child.iterchildren():
                if grand.tag.split("}")[-1] == "r":
                    link_text_parts.append(_runs_text(grand))

            link_text = "".join(link_text_parts).strip()
            if link_text:
                parts.append(
                    '<a href="{href}">{text}</a>'.format(
                        href=html.escape(url, quote=True),
                        text=html_entities(link_text)
                    )
                )
            continue

        # (B) fldSimple (campo semplice)
        if tag == "fldSimple":
            instr = child.get(qn("w:instr")) or child.get("{%s}instr" % W_NS) or ""
            url = _extract_url_from_instr(instr)

            display = []
            for node in child.iter():
                if node.tag.endswith("}t") and node.text:
                    display.append(node.text)
            visible = "".join(display).strip()

            if visible:
                parts.append(
                    '<a href="{href}">{text}</a>'.format(
                        href=html.escape(url, quote=True),
                        text=html_entities(visible)
                    )
                )
            continue

        # (C) Run: può contenere campi complessi
        if tag == "r":
            fldChar = child.find(".//w:fldChar", NS)
            if fldChar is not None:
                fld_type = fldChar.get(qn("w:fldCharType")) or fldChar.get("{%s}fldCharType" % W_NS)

                if fld_type == "begin":
                    if in_field:
                        flush_field()
                    in_field = True
                    field_instr = ""
                    field_url = None
                    after_separate = False
                    field_display_parts = []
                    continue

                if fld_type == "separate" and in_field:
                    field_url = _extract_url_from_instr(field_instr)
                    after_separate = True
                    continue

                if fld_type == "end" and in_field:
                    flush_field()
                    continue

            instrText = child.find(".//w:instrText", NS)
            if instrText is not None and in_field and not after_separate:
                if instrText.text:
                    field_instr += instrText.text
                continue

            txt = _runs_text(child)
            if not txt:
                continue

            if in_field:
                if after_separate:
                    field_display_parts.append(txt)
            else:
                parts.append(html_entities(txt))
            continue

    if in_field:
        flush_field()

    return "".join(parts).strip()


def extract_all_lines_as_html(doc_obj) -> List[str]:
    """
    Estrae righe da paragrafi e tabelle mantenendo l’ordine.
    Ogni riga è già “HTML-safe” e preserva i link come <a href="...">.

    MODIFICA #1:
    - Liste puntate/num. (paragrafi con w:numPr) vengono raggruppate e convertite in <ul><li>...</li></ul>.
    """
    lines: List[str] = []

    def add_line(t: str):
        t = (t or "").strip()
        if t:
            lines.append(t)

    def process_container(container):
        blocks = list(iter_block_items(container))
        i = 0
        while i < len(blocks):
            block = blocks[i]

            if isinstance(block, Paragraph):
                # ---- LISTE: raggruppa paragrafi consecutivi di lista in un <ul> ----
                if is_list_paragraph(block):
                    items: List[str] = []
                    while i < len(blocks) and isinstance(blocks[i], Paragraph) and is_list_paragraph(blocks[i]):
                        item_html = paragraph_to_html(blocks[i]).strip()
                        if item_html:
                            items.append(item_html)
                        i += 1

                    if items:
                        ul = "<ul>\n{}\n</ul>".format(
                            "\n".join(f"<li>{it}</li>" for it in items)
                        )
                        add_line(ul)
                    continue

                # ---- PARAGRAFO NORMALE ----
                add_line(paragraph_to_html(block))
                i += 1
                continue

            if isinstance(block, Table):
                # scendi nelle celle, preservando l’ordine
                for row in block.rows:
                    for cell in row.cells:
                        process_container(cell)
                i += 1
                continue

            i += 1

    process_container(doc_obj)
    return lines


# =========================
# Parsing input
# =========================

def parse_input_docx(path: Path) -> Dict[str, Any]:
    src = Document(str(path))
    lines = extract_all_lines_as_html(src)

    meta: Dict[str, str] = {k: "" for k in OUTPUT_META_LABELS}
    h1: str = ""
    testo_lines: List[str] = []

    in_testo = False
    saw_any_meta = False

    testo_re = re.compile(r"^({})\s*:\s*$".format("|".join(TESTO_KEYS)), re.IGNORECASE)

    for line in lines:
        if testo_re.match(line):
            in_testo = True
            continue

        if not in_testo:
            m = re.match(r"^([^:]{1,80})\s*:\s*(.*)$", line)
            if m:
                k_raw = _normalize_key(m.group(1))
                v = (m.group(2) or "").strip()
                if k_raw in INPUT_KEY_MAP:
                    out_k = INPUT_KEY_MAP[k_raw]
                    saw_any_meta = True
                    if out_k == "H1":
                        h1 = v
                    elif out_k in meta:
                        meta[out_k] = v
            continue

        testo_lines.append(line)

    # fallback: se "Testo:" manca
    if not testo_lines:
        non_meta = []
        for line in lines:
            m = re.match(r"^([^:]{1,80})\s*:\s*(.*)$", line)
            if m:
                k_raw = _normalize_key(m.group(1))
                if k_raw in INPUT_KEY_MAP or k_raw in TESTO_KEYS:
                    continue
            non_meta.append(line)
        testo_lines = non_meta

    # fallback H1
    if not h1:
        h1 = (meta.get("Title") or "").strip()
    if not h1 and testo_lines:
        h1 = testo_lines[0].strip()
    if not h1:
        h1 = "Untitled"

    # Parsing corpo: intro + sezioni (h2/h3)
    intro_paras: List[str] = []
    sections: List[Dict[str, Any]] = []
    current_title: Optional[str] = None
    current_level: str = "h2"
    current_paras: List[str] = []

    def flush_section():
        nonlocal current_title, current_level, current_paras, sections
        if current_title is not None:
            sections.append({"title": current_title, "level": current_level, "paras": current_paras[:]})
        current_title = None
        current_level = "h2"
        current_paras = []

    for t in testo_lines:
        t = (t or "").strip()
        if not t:
            continue

        # MODIFICA #2: distinguere h2 vs h3 già qui
        m = re.match(r"^(.*)\s*\((h2|h3)\)\s*$", t, re.IGNORECASE)
        if m:
            flush_section()
            current_title = m.group(1).strip()
            current_level = m.group(2).lower()
            continue

        if current_title is None:
            intro_paras.append(t)
        else:
            current_paras.append(t)

    flush_section()

    return {
        "meta": meta,
        "h1": h1.strip(),
        "intro_paras": [p for p in intro_paras if p],
        "sections": sections,
        "_debug": {
            "total_lines_extracted": len(lines),
            "testo_lines": len(testo_lines),
            "saw_any_meta": saw_any_meta,
        }
    }


# =========================
# HTML generation (final blocks)
# =========================

def _wrap_paragraph_or_passthrough(html_line: str) -> str:
    """
    Se la riga è un blocco HTML (es. <ul>...</ul>), la lascia così.
    Altrimenti la wrappa in <p class="...">...</p>.
    """
    t = (html_line or "").strip()
    if not t:
        return '<p class="h-text-size-14 h-font-primary"></p>'
    if t.startswith("<ul") or t.startswith("<ol"):
        return t
    return '<p class="h-text-size-14 h-font-primary">{}</p>'.format(t)

def build_html_rows(
    parsed: Dict[str, Any],
    n_images: int = 0,
    product_ids: Optional[List[str]] = None,
) -> List[Tuple[str, str]]:
    """
    Costruisce le righe Block/HTML allineate al formato di riferimento:

      ("H1",           "<h1>...</h1>")
      ("Intro",        "<p ...>...</p>")
      ("S1 IMAGE",     "")          ← cella vuota PRIMA della sezione
      ("S1",           "<h2>...</h2>  <p>...</p>")
      ("S2 IMAGE",     "")
      ("S2",           "...")
      ...
      ("➡️Related Product", "00266-LAC ; 01123-LAC")  ← ID grezzi, non HTML

    Le immagini vengono assegnate alle prime N sezioni (una per sezione),
    inserendo il placeholder IMAGE vuoto immediatamente prima della sezione.
    Le sezioni oltre n_images non ricevono IMAGE.
    """
    rows: List[Tuple[str, str]] = []
    if product_ids is None:
        product_ids = []

    # --- H1 ---
    h1 = parsed.get("h1") or ""
    rows.append(("H1", "<h1>{}</h1>".format(h1)))

    # --- Intro ---
    intro_paras = parsed.get("intro_paras") or []
    if intro_paras:
        intro_html = "\n\n".join(_wrap_paragraph_or_passthrough(p) for p in intro_paras)
    else:
        intro_html = '<p class="h-text-size-14 h-font-primary"></p>'
    rows.append(("Intro", intro_html))

    # --- Sezioni numerate con IMAGE prima (se prevista) ---
    for sec_idx, sec in enumerate(parsed.get("sections", [])):
        sec_num = sec_idx + 1
        title   = sec.get("title", "")
        level   = (sec.get("level") or "h2").lower()
        paras   = sec.get("paras", [])

        # Placeholder vuoto prima della sezione, se rientra nel conteggio immagini
        if sec_num <= n_images:
            rows.append(("S{} IMAGE".format(sec_num), ""))

        # Contenuto sezione
        if level == "h3":
            heading = '<h3><strong>{}</strong></h3>'.format(title)
        else:
            heading = '<h2><strong>{}</strong></h2>'.format(title)

        parts = [heading]
        parts.extend(_wrap_paragraph_or_passthrough(p) for p in paras if (p or "").strip())
        rows.append(("S{}".format(sec_num), "\n\n".join(parts).strip()))

    # --- Related Product (ID grezzi separati da " ; ") ---
    if product_ids:
        rows.append(("➡️Related Product", " ; ".join(product_ids)))

    return rows


def build_structure_of_content(html_rows: List[Tuple[str, str]]) -> List[str]:
    """
    Produce la lista per 'Structure of content'.
    Mostra H1, Intro e le sezioni Sx come '✏️ S3'.
    I blocchi IMAGE e Related Product non compaiono nella struttura.
    """
    s = []
    for block, _ in html_rows:
        if block == "H1":
            s.append("H1")
        elif block == "Intro":
            s.append("Intro")
        elif block.startswith("S") and not block.endswith("IMAGE") and block != "Intro":
            s.append("✏️ S3")
    return s


# =========================
# DOCX writer
# =========================

def write_output_docx(
    parsed: Dict[str, Any],
    output_path: Path,
    n_images: int = 0,
    product_ids: Optional[List[str]] = None,
):
    doc = Document()

    meta = parsed.get("meta", {})
    title = (meta.get("Title") or "").strip() or (parsed.get("h1") or "").strip() or "Untitled"

    # Titolo top
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("")

    # Tabella metadati
    table = doc.add_table(rows=len(OUTPUT_META_LABELS), cols=2)
    table.style = "Table Grid"

    for i, key in enumerate(OUTPUT_META_LABELS):
        left = table.cell(i, 0)
        right = table.cell(i, 1)

        shade_cell(left, "000000")
        set_cell_text(left, key, bold=True, color=RGBColor(255, 255, 255), size_pt=10)
        set_cell_text(right, (meta.get(key, "") or ""), bold=False, size_pt=10)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Structure of content
    p = doc.add_paragraph("Structure of content:")
    if p.runs:
        p.runs[0].bold = True
    else:
        p.add_run("Structure of content:").bold = True

    html_rows = build_html_rows(parsed, n_images=n_images, product_ids=product_ids or [])
    for line in build_structure_of_content(html_rows):
        doc.add_paragraph(line)

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Tabella Block | HTML Output
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"

    hdr0 = t2.cell(0, 0)
    hdr1 = t2.cell(0, 1)
    shade_cell(hdr0, "D9D9D9")
    shade_cell(hdr1, "D9D9D9")
    set_cell_text(hdr0, "Block", bold=True, size_pt=10)
    set_cell_text(hdr1, "⭐ HTML Output ⭐", bold=True, size_pt=10)

    for block, html_block in html_rows:
        row_cells = t2.add_row().cells
        row_cells[0].text = block
        row_cells[1].text = html_block

        # evidenzia visivamente le righe speciali
        if "IMAGE" in block:
            shade_cell(row_cells[0], "FFF2CC")  # giallo chiaro
            shade_cell(row_cells[1], "FFF2CC")
        elif "Related Product" in block:
            shade_cell(row_cells[0], "D9EAD3")  # verde chiaro
            shade_cell(row_cells[1], "D9EAD3")

        for cell in row_cells:
            for para in cell.paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)

    doc.save(str(output_path))


# =========================
# Runner
# =========================

def convert_one(input_docx: Path, output_docx: Path, ped_xlsx: Optional[Path] = None):
    parsed = parse_input_docx(input_docx)

    dbg = parsed.get("_debug", {})
    print("DEBUG:", dbg)
    print("DEBUG meta:", parsed.get("meta", {}))
    print("DEBUG h1:", parsed.get("h1", ""))
    print("DEBUG intro_paras:", len(parsed.get("intro_paras", [])))
    print("DEBUG sections:", len(parsed.get("sections", [])))

    # Lookup Excel per immagini e carosello
    extras = {"n_images": 0, "product_ids": []}
    if ped_xlsx:
        lookup = load_ped_lookup(ped_xlsx)
        slug = (parsed.get("meta", {}).get("URL") or "").strip()
        h1   = (parsed.get("h1") or "").strip()
        extras = lookup_article_extras(lookup, slug=slug, h1=h1)
        print("DEBUG extras (PED lookup):", extras)

    write_output_docx(
        parsed,
        output_docx,
        n_images=extras["n_images"],
        product_ids=extras["product_ids"],
    )

if __name__ == "__main__":
    inp = Path("NVL_202601_Lancome_Ottimizzazione-Pagina-Ingrediente-Collagene-v2.docx")
    out = Path("output_NVL_202601_Lancome_Ottimizzazione-Pagina-Ingrediente-Collagene-v2.docx")
    ped = Path("202603_LAN_PED-2026.xlsx")  # opzionale: se non esiste, i placeholder vengono saltati
    convert_one(inp, out, ped_xlsx=ped if ped.exists() else None)
    print("Creato:", out)