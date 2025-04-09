import streamlit as st
import pandas as pd
import openai
import os
from docx import Document
from docx.shared import Pt
import tempfile
import zipfile

# 📌 Titolo della Web App
st.title("OpenAI Content Generation 📝")

# 📌 Inserimento della chiave API di OpenAI
def get_openai_api_key():
    with st.sidebar:
        st.subheader("🔐 API Key OpenAI")
        if 'oai_api_key' not in st.session_state:
            st.session_state['oai_api_key'] = ""

        oai_api_key = st.text_input(
            "Insert your OpenAI API KEY",
            type="password",
            value=st.session_state['oai_api_key']
        )

        if not oai_api_key:
            st.warning("⚠️ Insert your API KEY to continue.")
            st.stop()

        # Salva in sessione
        st.session_state['oai_api_key'] = oai_api_key
        return oai_api_key
    oai_client = openai.Client(api_key=oai_api_key)

# 📌 Caricamento del file Excel
uploaded_file = st.file_uploader("Upload Excel file", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    df.columns = df.columns.str.strip()

    # 📌 Controllo se esistono le colonne "Prompt System" e "Numero Minimo Parole"
    if "Prompt System" not in df.columns:
        df["Prompt System"] = "Sei un copywriter esperto in SEO."
    
    if "Numero Minimo Parole" not in df.columns:
        df["Numero Minimo Parole"] = 800  # Valore predefinito se la colonna non esiste

    st.write("📊 File Preview:")
    st.dataframe(df.head())

    # 📌 Crea una cartella temporanea per i file generati
    temp_dir = tempfile.mkdtemp()

    # 📌 Funzione per contare le parole in un testo
    def count_words(text):
        return len(text.split())

    # ✨ Generazione di un singolo paragrafo
    def generate_paragraph(title, description, system_prompt):
        prompt = f"""
        Scrivi un paragrafo dettagliato con il seguente titolo e descrizione:
        **Titolo:** {title}
        **Descrizione:** {description}
        """

        response = oai_client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "system", "content": system_prompt},
                      {"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000
        )

        return response.choices[0].message.content

    # ✨ Generazione del contenuto totale con controllo sulla lunghezza
    def generate_full_content(row):
        min_words = row["Numero Minimo Parole"]  # Ora sempre presente grazie al controllo iniziale
        min_threshold = int(min_words * 0.80)  # 80% del minimo richiesto
        system_prompt = row["Prompt System"]  # Usa il prompt personalizzato per ogni riga

        attempts = 0
        max_attempts = 5
        best_text = ""
        best_word_count = 0

        while attempts < max_attempts:
            full_text = []
            for i in range(1, 11):
                title_col = f"Paragrafo {i} (Titolo)"
                desc_col = f"Paragrafo {i} (Descrizione)"
                if title_col in row and pd.notna(row[title_col]) and pd.notna(row[desc_col]):
                    paragraph_text = generate_paragraph(row[title_col], row[desc_col], system_prompt)
                    full_text.append(f"{row[title_col]}\n\n{paragraph_text}")

            full_text_str = "\n\n".join(full_text)
            word_count = count_words(full_text_str)

            if word_count >= min_threshold:
                return full_text_str  # ✅ Testo valido, restituito subito

            # 🔄 Salviamo il miglior risultato nel caso nessun tentativo superi la soglia
            if word_count > best_word_count:
                best_text = full_text_str
                best_word_count = word_count

            attempts += 1

        return best_text  # 🔄 Restituisce il miglior risultato trovato

    # ✨ Creazione del file Word con formattazione avanzata
    def create_word_document(row, generated_text):
        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run(f"SEO Title: {row['SEO Title']}")
        run.bold = True
        run.font.size = Pt(14)
        meta = doc.add_paragraph()
        run = meta.add_run(f"Meta Description: {row['Meta Description']}")
        run.italic = True
        run.font.size = Pt(12)

        for paragraph in generated_text.split("\n"):
            if paragraph.strip() == "":
                continue
            if paragraph.endswith(":"):
                p = doc.add_paragraph()
                p.style = "Heading 2"
                run = p.add_run(paragraph)
                run.bold = True
                run.font.size = Pt(14)
            elif paragraph.startswith("-"):
                doc.add_paragraph(paragraph, style="List Bullet")
            else:
                doc.add_paragraph(paragraph)

        file_path = os.path.join(temp_dir, f"{row['Titolo'].replace(' ', '_')}.docx")
        doc.save(file_path)
        return file_path

    if st.button("Contents Generation"):
        generated_files = []
        for index, row in df.iterrows():
            st.write(f"📝 Generating content for: {row['Titolo']}...")
            full_text = generate_full_content(row)
            if full_text:
                file_path = create_word_document(row, full_text)
                generated_files.append(file_path)

        zip_path = os.path.join(temp_dir, "contenuti_generati.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file in generated_files:
                zipf.write(file, os.path.basename(file))

        with open(zip_path, "rb") as f:
            st.download_button("📥 Download contents.", f, file_name="contenuti_generati.zip")