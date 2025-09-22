import streamlit as st
from openai import OpenAI
from youtube_transcript_api import YouTubeTranscriptApi as YTA, TranscriptsDisabled, NoTranscriptFound
import re
from docx import Document
import io
import zipfile
from datetime import datetime

# ============================
# Helpers
# ============================

def fetch_transcript(video_id: str, language_code: str, allow_fallback: bool = True):
    """Robust transcript fetcher that works with youtube-transcript-api list_transcripts API.
    - Tries preferred language
    - If allowed, tries to translate available transcripts to the requested language
    - Finally, falls back to the first available transcript
    """
    transcripts = YTA.list_transcripts(video_id)

    # 1) Try exact language match
    try:
        return transcripts.find_transcript([language_code]).fetch()
    except NoTranscriptFound:
        if not allow_fallback:
            raise

    # 2) Try translating any available transcript to the requested language
    if allow_fallback:
        for tr in transcripts:
            if getattr(tr, "is_translatable", False):
                try:
                    return tr.translate(language_code).fetch()
                except Exception:
                    pass

    # 3) Last resort: return the first available transcript as-is
    for tr in transcripts:
        try:
            return tr.fetch()
        except Exception:
            pass

    # If nothing worked
    raise NoTranscriptFound(f"No transcript could be retrieved for video {video_id}.")

# ============================
# Helpers
# ============================

def extract_video_id(url: str):
    """Extract the 11-char YouTube video ID from many URL formats."""
    if not url:
        return None
    # Common patterns: v=, youtu.be/, /embed/, /shorts/
    patterns = [
        r"(?:v=|/v/|&v=|watch\?v=)([a-zA-Z0-9_-]{11})",
        r"youtu\.be/([a-zA-Z0-9_-]{11})",
        r"/embed/([a-zA-Z0-9_-]{11})",
        r"/shorts/([a-zA-Z0-9_-]{11})",
    ]
    for p in patterns:
        m = re.search(p, url)
        if m:
            return m.group(1)
    return None


def clean_transcript_text(text: str) -> str:
    """Basic cleanup for transcript artifacts."""
    text = re.sub(r"\[(?:Music|Applause|Laughter)\]", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def summarize_text(client: OpenAI, text: str, model: str, temperature: float, min_words: int = 400) -> str:
    """Summarize arbitrarily long text by chunking if needed."""
    # Rough chunking by characters to avoid context limits
    # Target ~12k chars per chunk; adjust if needed per model
    MAX_CHARS = 12000
    chunks = [text[i:i + MAX_CHARS] for i in range(0, len(text), MAX_CHARS)]

    if len(chunks) == 1:
        prompt = (
            f"Summarize the following transcript, keeping the key points and context. "
            f"Ensure the summary is at least {min_words} words, clear, concise, non-repetitive, and highlights the most important aspects.\n\n"
            f"Transcript:\n{chunks[0]}\n\n"
            "Rules:\n- Keep it clear and structured (use short paragraphs or bullet points if helpful).\n"
            "- Avoid redundancy and trivial details.\n- Preserve all critical explanations, definitions, and conclusions."
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert at synthesizing video transcripts into accurate, readable long summaries."},
                {"role": "user", "content": prompt},
            ],
            temperature=temperature,
            max_tokens=4000,
        )
        return resp.choices[0].message.content

    # Multi-step: summarize chunks then synthesize
    partial_summaries = []
    for idx, ch in enumerate(chunks, 1):
        prompt = (
            f"Summarize part {idx} of a longer transcript. Focus on key points, arguments, data, and conclusions. "
            f"Length: ~250-350 words.\n\nPart {idx}:\n{ch}"
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert note-taker."},
                {"role": "user", "content": prompt},
            ],
            temperature=temperature,
            max_tokens=1200,
        )
        partial_summaries.append(resp.choices[0].message.content)

    synthesis_prompt = (
        f"You are given {len(partial_summaries)} partial summaries from a video transcript. "
        f"Merge them into a single cohesive summary of at least {min_words} words. "
        "Avoid duplication, keep a logical flow, and highlight the most critical insights."
        "\n\nPartial summaries:\n" + "\n\n".join(f"Part {i+1}:\n{ps}" for i, ps in enumerate(partial_summaries))
    )
    final_resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are an expert editor."},
            {"role": "user", "content": synthesis_prompt},
        ],
        temperature=temperature,
        max_tokens=4000,
    )
    return final_resp.choices[0].message.content


# ============================
# UI
# ============================

st.set_page_config(page_title="YouTube Transcript & Summary Generator", page_icon="🎥", layout="centered")
st.title("🎥 YouTube Transcript & Summary Generator")

# --- Sidebar: OpenAI settings ---
with st.sidebar:
    st.subheader("🔐 OpenAI Settings")
    if 'oai_api_key' not in st.session_state:
        st.session_state['oai_api_key'] = ""

    oai_key = st.text_input("Insert your OpenAI API KEY", type="password", value=st.session_state['oai_api_key'])
    st.session_state['oai_api_key'] = oai_key

    model = st.selectbox(
        "Model",
        # Keep v1-compatible names; user can switch if needed
        options=[
            "gpt-4o-mini",
            "gpt-4o",
            "gpt-4.1-mini",
            "gpt-4.1",
            "gpt-4",
        ],
        index=0,
        help="Choose an OpenAI model. All listed are compatible with the v1 Python SDK's chat.completions API."
    )
    temperature = st.slider("Creativity (temperature)", 0.0, 1.0, 0.7, 0.1)
    min_words = st.number_input("Minimum summary length (words)", min_value=150, max_value=2000, value=400, step=50)

# --- Main inputs ---
youtube_url = st.text_input("📺 Enter YouTube video URL")
language_code = st.text_input("🌍 Enter language code (default 'en')", value="en")
allow_fallback = st.checkbox("If not found, try other available languages", value=True)

# --- How to use (English guide) ---
with st.expander("🧭 How to use (English)"):
    st.markdown(
        """
        1. **Enter your OpenAI API key** in the sidebar (required). You can create one in your OpenAI dashboard.
        2. **Paste a YouTube video URL** into the input field above.
        3. Optionally **set the transcript language code** (e.g., `en`, `it`, `es`).
        4. Adjust **model**, **creativity**, and **minimum summary length** in the sidebar if needed.
        5. Click **"Generate Transcript and Summary"**.
        6. When finished, **download the ZIP** containing a `.txt` transcript and a `.docx` summary.
        
        **Notes**
        - If the video has **no transcript** or **transcripts are disabled**, you'll see a clear error.
        - If the requested language isn't available and fallback is enabled, the app tries other languages.
        - Very long videos are handled by **chunked summarization** to keep results stable.
        """
    )

# Guard clause for API key
if not oai_key:
    st.warning("⚠️ Insert your API KEY in the sidebar to continue.")
    st.stop()

client = OpenAI(api_key=oai_key)

# --- Action button ---
if st.button("📄 Generate Transcript and Summary", use_container_width=True):
    if not youtube_url:
        st.error("⚠️ You must enter the video URL!")
        st.stop()

    video_id = extract_video_id(youtube_url)
    if not video_id:
        st.error("❌ Invalid YouTube URL. Please try again.")
        st.stop()

    try:
        # 1) Fetch transcript
        try:
            transcript = fetch_transcript(video_id, language_code, allow_fallback=allow_fallback)
        except NoTranscriptFound:
            if allow_fallback:
                st.warning(f"⚠️ No transcript found in '{language_code}'. Trying other available languages...")
                transcript = fetch_transcript(video_id, language_code, allow_fallback=True)
            else:
                raise
        except TranscriptsDisabled:
            st.error("❌ Transcripts are disabled for this video.")
            st.stop()

        if not transcript:
            st.error("❌ No transcript available for this video.")
            st.stop()

        full_text = " ".join([entry.get('text', '') for entry in transcript])
        full_text = clean_transcript_text(full_text)

        # 2) Summarize with OpenAI
        with st.spinner("Generating summary with OpenAI..."):
            summary = summarize_text(client, full_text, model=model, temperature=temperature, min_words=int(min_words))

        # 3) Build in-memory files
        # 3a) Transcript .txt
        transcript_bytes = full_text.encode('utf-8')
        transcript_filename = f"transcript_{video_id}.txt"

        # 3b) Summary .docx (in-memory)
        doc = Document()
        doc.add_heading("Riassunto del Video", level=1)
        doc.add_paragraph(summary)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        summary_filename = f"summary_{video_id}.docx"

        # 3c) ZIP both
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(transcript_filename, transcript_bytes)
            zf.writestr(summary_filename, docx_buffer.getvalue())
        zip_buffer.seek(0)
        zip_filename = f"youtube_summary_{video_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

        st.success("✅ Transcript and Summary successfully generated!")
        st.download_button(
            "📥 Download ZIP with Transcript & Summary",
            data=zip_buffer,
            file_name=zip_filename,
            mime="application/zip",
            use_container_width=True,
        )

        # Also show a preview
        with st.expander("👀 Preview summary"):
            st.write(summary)

    except NoTranscriptFound:
        st.error(f"❌ No transcript found for this video (language: '{language_code}'). Try enabling fallback or choosing a different language.")
    except Exception as e:
        st.error(f"❌ Error during extraction: {str(e)}")