import streamlit as st
from openai import OpenAI
import re
from urllib.parse import urlparse, parse_qs
from docx import Document
import io
import zipfile
from datetime import datetime
import time
import importlib
import json
import xml.etree.ElementTree as ET
import math
import string

# ============================
# Eccezione interna cross-version
# ============================
class NoTranscriptAvailable(Exception):
    """Eccezione interna: nessuna trascrizione disponibile via API/yt_dlp."""

# ============================
# Import youtube-transcript-api dinamico + introspezione
# ============================
YTA = None
TranscriptsDisabled = Exception
NoTranscriptFound = Exception
YTA_VERSION = "unknown"
YTA_FILE = "unknown"

try:
    yta_mod = importlib.import_module("youtube_transcript_api")
    YTA_FILE = getattr(yta_mod, "__file__", "unknown")
    YTA_VERSION = getattr(yta_mod, "__version__", "unknown")
    YTA = getattr(yta_mod, "YouTubeTranscriptApi", None)
    TranscriptsDisabled = getattr(yta_mod, "TranscriptsDisabled", Exception)
    NoTranscriptFound = getattr(yta_mod, "NoTranscriptFound", Exception)
except Exception as e:
    st.error(f"❌ Impossibile importare 'youtube_transcript_api': {type(e).__name__}: {e}")
    st.stop()

# ============================
# Helpers: estrazione ID e retry
# ============================
def extract_video_id(url: str):
    """Estrae l'ID (11 char) da molte varianti di URL YouTube."""
    if not url:
        return None
    try:
        p = urlparse(url)
        if p.netloc:
            qs = parse_qs(p.query)
            if 'v' in qs and len(qs['v'][0]) == 11:
                return qs['v'][0]
            m = re.search(r"/shorts/([A-Za-z0-9_-]{11})", p.path)
            if m: return m.group(1)
            m = re.search(r"/embed/([A-Za-z0-9_-]{11})", p.path)
            if m: return m.group(1)
            m = re.search(r"youtu\.be/([A-Za-z0-9_-]{11})", p.netloc + p.path)
            if m: return m.group(1)
    except Exception:
        pass
    for pat in [
        r"(?:v=|/v/|&v=|watch\?v=)([A-Za-z0-9_-]{11})",
        r"youtu\.be/([A-Za-z0-9_-]{11})",
        r"/embed/([A-Za-z0-9_-]{11})",
        r"/shorts/([A-Za-z0-9_-]{11})",
    ]:
        m = re.search(pat, url)
        if m: return m.group(1)
    return None

def _with_retry(fn, *args, max_attempts=4, **kwargs):
    """Retry con backoff esponenziale (0.5→1→2s) per errori temporanei."""
    delay = 0.5
    for attempt in range(max_attempts):
        try:
            return fn(*args, **kwargs)
        except (TranscriptsDisabled, NoTranscriptFound, NoTranscriptAvailable):
            raise
        except Exception:
            if attempt == max_attempts - 1:
                raise
            time.sleep(delay)
            delay = min(2.0, delay * 2)

# ============================
# Parsers per sottotitoli in formati diversi
# ============================
def _parse_time_to_seconds(s: str) -> float:
    parts = s.split(":")
    parts = [p.replace(",", ".") for p in parts]
    while len(parts) < 3:
        parts = ["0"] + parts
    h, m, sec = float(parts[-3]), float(parts[-2]), float(parts[-1])
    return h*3600 + m*60 + sec

def parse_vtt(text: str):
    """Parsa VTT in entries [{'text','start','duration'}] provando webvtt se disponibile, altrimenti regex fallback."""
    entries = []
    try:
        import webvtt  # webvtt-py
        v = webvtt.read_buffer(io.StringIO(text))
        for cue in v:
            start = _parse_time_to_seconds(cue.start)
            end = _parse_time_to_seconds(cue.end)
            duration = max(0.0, end - start)
            txt = re.sub(r"\s+", " ", (cue.text or "").strip())
            if txt:
                entries.append({"text": txt, "start": start, "duration": duration})
        if entries:
            return entries
    except Exception:
        pass
    # fallback minimale
    for block in re.split(r"\n\s*\n", text.strip()):
        lines = [ln.strip() for ln in block.splitlines() if ln.strip()]
        if len(lines) < 2:
            continue
        time_line = next((ln for ln in lines if "-->" in ln), None)
        if not time_line:
            continue
        tm = re.match(r"(?P<start>[\d:.]+)\s*-->\s*(?P<end>[\d:.]+)", time_line)
        if not tm:
            continue
        start = _parse_time_to_seconds(tm.group("start"))
        end = _parse_time_to_seconds(tm.group("end"))
        duration = max(0.0, end - start)
        text_lines = []
        for ln in lines:
            if ln == time_line: 
                continue
            if re.fullmatch(r"\d+", ln):
                continue
            text_lines.append(ln)
        txt = re.sub(r"\s+", " ", " ".join(text_lines).strip())
        if txt:
            entries.append({"text": txt, "start": start, "duration": duration})
    return entries

def parse_json3(text: str):
    data = json.loads(text)
    events = data.get("events", [])
    entries = []
    for ev in events:
        t_start = ev.get("tStartMs")
        dur = ev.get("dDurationMs", 0)
        segs = ev.get("segs", [])
        if t_start is None or not segs:
            continue
        txt = " ".join(seg.get("utf8", "") for seg in segs).strip()
        txt = re.sub(r"\s+", " ", txt)
        if not txt:
            continue
        start = float(t_start) / 1000.0
        duration = float(dur) / 1000.0 if dur else 0.0
        entries.append({"text": txt, "start": start, "duration": duration})
    return entries

def parse_srv3_xml(text: str):
    entries = []
    root = ET.fromstring(text)
    for p in root.iter("p"):
        t = p.attrib.get("t")  # ms
        d = p.attrib.get("d", "0")
        if t is None:
            continue
        start = float(t) / 1000.0
        duration = float(d) / 1000.0 if d else 0.0
        parts = []
        if p.text and p.text.strip():
            parts.append(p.text.strip())
        for s in p.iter("s"):
            if s.text and s.text.strip():
                parts.append(s.text.strip())
        txt = re.sub(r"\s+", " ", " ".join(parts).strip())
        if txt:
            entries.append({"text": txt, "start": start, "duration": duration})
    return entries

def parse_ttml_xml(text: str):
    entries = []
    root = ET.fromstring(text)
    for p in root.iter():
        if p.tag.endswith("}p") or p.tag == "p":
            begin = p.attrib.get("begin")
            end = p.attrib.get("end")
            if not begin or not end:
                continue
            start = _parse_time_to_seconds(begin)
            end_s = _parse_time_to_seconds(end)
            duration = max(0.0, end_s - start)
            txt = "".join(p.itertext()).strip()
            txt = re.sub(r"\s+", " ", txt)
            if txt:
                entries.append({"text": txt, "start": start, "duration": duration})
    return entries

# ============================
# Fallback via yt_dlp (esteso e robusto)
# ============================
def fetch_transcript_via_ytdlp(video_id: str, language_code: str, allow_fallback: bool, save_raw_debug: bool = False):
    """Prova tutte le tracce (manuali/auto) e formati (json3/srv3/ttml/vtt)."""
    try:
        import yt_dlp
        import requests
    except Exception:
        raise NoTranscriptAvailable("yt_dlp/requests non installati per il fallback.")

    url = f"https://www.youtube.com/watch?v={video_id}"
    ydl_opts = {"skip_download": True, "quiet": True, "no_warnings": True}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)

        def ordered_langs(subs_dict):
            langs = []
            if not subs_dict:
                return langs
            exact, base, others = [], [], []
            base_code = language_code.split("-")[0].lower()
            for lang in subs_dict.keys():
                if lang == language_code:
                    exact.append(lang)
                elif lang == base_code:
                    base.append(lang)
                else:
                    others.append(lang)
            return exact + base + (others if allow_fallback else [])

        candidates = []
        subs = info.get("subtitles") or {}
        for lang in ordered_langs(subs):
            for track in subs.get(lang, []):
                candidates.append(("subs", lang, track))
        auto = info.get("automatic_captions") or {}
        for lang in ordered_langs(auto):
            for track in auto.get(lang, []):
                candidates.append(("auto", lang, track))

        if not candidates:
            raise NoTranscriptAvailable("Nessun sottotitolo disponibile via yt_dlp.")

        last_error = None
        for label, lang, track in candidates:
            try:
                ext = (track.get("ext") or "").lower()
                track_url = track.get("url")
                if not track_url:
                    continue
                import requests
                r = _with_retry(lambda: requests.get(track_url, timeout=20))
                raw = r.text

                if save_raw_debug:
                    st.download_button(
                        f"⬇️ Scarica sottotitoli RAW ({label}:{lang}.{ext or 'txt'})",
                        data=raw.encode("utf-8", errors="ignore"),
                        file_name=f"{video_id}_{label}_{lang}.{ext or 'txt'}",
                        mime="text/plain"
                    )

                if ext == "json3":
                    entries = parse_json3(raw)
                elif ext in ("srv3", "xml"):
                    try:
                        entries = parse_srv3_xml(raw)
                        if not entries:
                            entries = parse_ttml_xml(raw)
                    except Exception:
                        entries = parse_ttml_xml(raw)
                elif ext in ("ttml",):
                    entries = parse_ttml_xml(raw)
                else:
                    entries = parse_vtt(raw)

                if entries:
                    return entries
                else:
                    last_error = f"Traccia {label}:{lang} ({ext}) non parsabile"
            except Exception as e:
                last_error = f"{label}:{lang} parsing error ({type(e).__name__}: {e})"
                continue

        raise NoTranscriptAvailable(last_error or "Impossibile parsare qualsiasi traccia sottotitoli via yt_dlp.")

# ============================
# Fetch transcript – percorsi modern/compat + orchestratore
# ============================
def fetch_transcript_modern(video_id: str, language_code: str, allow_fallback: bool = True):
    """Percorso moderno (quando esiste list_transcripts)."""
    base_lang = language_code.split("-")[0].lower() if language_code else "en"
    prefer_list = []
    if language_code:
        prefer_list.append(language_code)
    if base_lang and base_lang not in prefer_list:
        prefer_list.append(base_lang)

    if YTA and hasattr(YTA, "get_transcript"):
        try:
            return _with_retry(getattr(YTA, "get_transcript"), video_id, languages=prefer_list)
        except NoTranscriptFound:
            pass
        except Exception:
            pass

    transcripts = _with_retry(getattr(YTA, "list_transcripts"), video_id)

    try:
        return transcripts.find_transcript(prefer_list).fetch()
    except NoTranscriptFound:
        if not allow_fallback:
            raise NoTranscriptAvailable("Transcript non disponibile nella lingua richiesta.")

    if allow_fallback:
        for tr in transcripts:
            try:
                if getattr(tr, "is_translatable", False):
                    for tgt in prefer_list:
                        try:
                            return tr.translate(tgt).fetch()
                        except Exception:
                            continue
            except Exception:
                continue

    for tr in transcripts:
        try:
            return tr.fetch()
        except Exception:
            continue

    raise NoTranscriptAvailable("Nessuna trascrizione recuperabile (modern).")

def fetch_transcript_any(video_id: str, language_code: str, allow_fallback: bool = True, save_raw_debug: bool = False):
    """1) list_transcripts → 2) get_transcript → 3) yt_dlp."""
    base_lang = language_code.split("-")[0].lower() if language_code else "en"
    prefer_list = [lc for lc in [language_code, base_lang] if lc]

    if YTA and hasattr(YTA, "list_transcripts"):
        try:
            entries = fetch_transcript_modern(video_id, language_code, allow_fallback)
            return entries, "list_transcripts"
        except (TranscriptsDisabled, NoTranscriptFound, NoTranscriptAvailable):
            pass

    if YTA and hasattr(YTA, "get_transcript"):
        langs = prefer_list[:]
        if allow_fallback:
            for extra in ["en", "en-US", "en-GB", "it", "es", "de", "fr", "pt", "pt-BR"]:
                if extra not in langs:
                    langs.append(extra)
        for lc in langs:
            try:
                entries = _with_retry(getattr(YTA, "get_transcript"), video_id, languages=[lc])
                return entries, f"get_transcript({lc})"
            except (TranscriptsDisabled, NoTranscriptFound):
                continue
            except Exception:
                continue

    entries = fetch_transcript_via_ytdlp(video_id, language_code, allow_fallback, save_raw_debug=save_raw_debug)
    return entries, "yt_dlp"

# ============================
# Helpers: capitoli, pulizia, sintesi
# ============================
def get_chapters(video_id: str):
    """Recupera i capitoli via yt_dlp se disponibili."""
    try:
        import yt_dlp
        url = f"https://www.youtube.com/watch?v={video_id}"
        with yt_dlp.YoutubeDL({"skip_download": True, "quiet": True, "no_warnings": True}) as ydl:
            info = ydl.extract_info(url, download=False)
            return info.get("chapters") or []   # [{'start_time':sec, 'end_time':sec, 'title':...}, ...]
    except Exception:
        return []

def text_by_time_window(entries, start_s, end_s):
    return " ".join(e["text"] for e in entries if start_s <= e.get("start", 0) < end_s).strip()

def clean_transcript_text(text: str) -> str:
    text = re.sub(r"\[(?:Music|Applause|Laughter)\]", " ", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def summarize_text(client: OpenAI, text: str, model: str, temperature: float, min_words: int = 400) -> str:
    MAX_CHARS = 12000
    chunks = [text[i:i + MAX_CHARS] for i in range(0, len(text), MAX_CHARS)]
    if len(chunks) == 1:
        prompt = (
            f"Summarize the following transcript, keeping the key points and context. "
            f"Ensure the summary is at least {min_words} words, clear, concise, non-repetitive, and highlights the most important aspects.\n\n"
            f"Transcript:\n{chunks[0]}\n\n"
            "Rules:\n- Keep it clear and structured (use short paragraphs or bullet points if helpful).\n"
            "- Avoid redundancy and trivial details.\n- Preserve all critical explanations, definitions, and conclusions."
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert at synthesizing video transcripts into accurate, readable long summaries."},
                {"role": "user", "content": prompt},
            ],
            temperature=temperature,
            max_tokens=4000,
        )
        return resp.choices[0].message.content

    partial_summaries = []
    for idx, ch in enumerate(chunks, 1):
        prompt = (
            f"Summarize part {idx} of a longer transcript. Focus on key points, arguments, data, and conclusions. "
            f"Length: ~250-350 words.\n\nPart {idx}:\n{ch}"
        )
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are an expert note-taker."},
                {"role": "user", "content": prompt},
            ],
            temperature=temperature,
            max_tokens=1200,
        )
        partial_summaries.append(resp.choices[0].message.content)

    synthesis_prompt = (
        f"You are given {len(partial_summaries)} partial summaries from a video transcript. "
        f"Merge them into a single cohesive summary of at least {min_words} words. "
        "Avoid duplication, keep a logical flow, and highlight the most critical insights."
        "\n\nPartial summaries:\n" + "\n".join(f"Part {i+1}:\n{ps}" for i, ps in enumerate(partial_summaries))
    )
    final_resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are an expert editor."},
            {"role": "user", "content": synthesis_prompt},
        ],
        temperature=temperature,
        max_tokens=4000,
    )
    return final_resp.choices[0].message.content

STOPWORDS = set("""
a ad ai al allo aiu agli all agl alla alle con col coi da dal dallo dai dagli dalla dalle di del dello dei degli della delle
in nel nello nei negli nella nelle su sul sullo sui sugli sulla sulle per tra fra il lo la i gli le un uno una
che chi cui come dove quando perché perche quanto questa questo queste questi quell quello quella quegli quelle
""".split())

def extractive_summary(text: str, max_sentences: int = 20) -> str:
    sentences = re.split(r'(?<=[\.\?\!])\s+', text)
    if len(sentences) <= max_sentences:
        return text
    freq = {}
    translator = str.maketrans("", "", string.punctuation + "“”‘’«»…")
    for sent in sentences:
        for w in sent.lower().translate(translator).split():
            if w in STOPWORDS or len(w) <= 2:
                continue
            freq[w] = freq.get(w, 0) + 1
    if not freq:
        return " ".join(sentences[:max_sentences])
    max_f = max(freq.values())
    scores = []
    for i, sent in enumerate(sentences):
        score = 0.0
        words = sent.lower().translate(translator).split()
        for w in words:
            if w in freq:
                score += freq[w] / max_f
        if words:
            score /= math.log2(10 + len(words))
        scores.append((score, i, sent))
    top = sorted(scores, key=lambda x: x[0], reverse=True)[:max_sentences]
    top_sorted = [s for _, _, s in sorted(top, key=lambda x: x[1])]
    return " ".join(top_sorted)

def summarize_text_safe_mapreduce(client: OpenAI, text: str, model: str, temperature: float, min_words: int) -> str:
    """Map-Reduce con chunk più piccoli (token-safe)."""
    MAX_CHARS = 6000
    chunks = [text[i:i+MAX_CHARS] for i in range(0, len(text), MAX_CHARS)]
    partials = []
    for ch in chunks:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a concise note-taker."},
                {"role": "user", "content": f"Summarize in ~120-160 words:\n\n{ch}"}
            ],
            temperature=temperature,
            max_tokens=300,
        )
        partials.append(resp.choices[0].message.content)
    final = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are an expert editor."},
            {"role": "user", "content": "Merge these notes into a clear summary of at least "
                                        f"{min_words} words:\n\n" + "\n\n".join(partials)}
        ],
        temperature=temperature,
        max_tokens=1000,
    )
    return final.choices[0].message.content

# ============================
# UI
# ============================
st.set_page_config(page_title="YouTube Transcript & Summary Generator", page_icon="🎥", layout="centered")
st.title("🎥 YouTube Transcript & Summary Generator")

with st.sidebar:
    st.subheader("🔐 OpenAI Settings")
    if 'oai_api_key' not in st.session_state:
        st.session_state['oai_api_key'] = ""
    oai_key = st.text_input("Insert your OpenAI API KEY", type="password", value=st.session_state['oai_api_key'])
    st.session_state['oai_api_key'] = oai_key

    # Modelli (incluso GPT-5)
    model = st.selectbox(
        "Model",
        options=[
            "gpt-5",
            "gpt-5-chat-latest",
            "gpt-5-mini",
            "gpt-4o",
            "gpt-4o-mini",
            "gpt-4.1",
            "gpt-4.1-mini",
            "gpt-4",
        ],
        index=0,
        help="Scegli un modello compatibile con l'API chat.completions."
    )
    temperature = st.slider("Creativity (temperature)", 0.0, 1.0, 0.7, 0.1)
    min_words = st.number_input("Minimum summary length (words)", min_value=150, max_value=2000, value=400, step=50)

# Diagnostica della libreria caricata
st.caption(f"📦 youtube-transcript-api — version: {YTA_VERSION}, file: {YTA_FILE}")
st.caption(f"🔍 capabilities → list_transcripts: {hasattr(YTA, 'list_transcripts') if YTA else False}, get_transcript: {hasattr(YTA, 'get_transcript') if YTA else False}")

youtube_url = st.text_input("📺 Enter YouTube video URL")
language_code = st.text_input("🌍 Enter language code (default 'en')", value="en")
allow_fallback = st.checkbox("If not found, try other available languages", value=True)
save_raw_debug = st.checkbox("🐞 Debug: salva sottotitoli raw non parsati", value=False)

# Modalità di sintesi
summary_mode = st.selectbox(
    "🧠 Summarization mode",
    [
        "Standard (auto-chunk)",
        "Map-Reduce (token-safe)",
        "Chapter-based Map-Reduce (if available)",
        "Extractive (no OpenAI)"
    ],
    index=0
)

if not oai_key:
    st.warning("⚠️ Insert your API KEY in the sidebar to continue.")
    st.stop()

client = OpenAI(api_key=oai_key)

if st.button("📄 Generate Transcript and Summary", use_container_width=True):
    if not youtube_url:
        st.error("⚠️ You must enter the video URL!")
        st.stop()

    video_id = extract_video_id(youtube_url)
    if not video_id:
        st.error("❌ Invalid YouTube URL: unable to extract a valid video ID.")
        st.stop()

    st.caption(f"🎯 Detected Video ID: `{video_id}`")

    try:
        # 1) Transcript (modern -> compat -> yt_dlp)
        entries, used_path = fetch_transcript_any(video_id, language_code, allow_fallback=allow_fallback, save_raw_debug=save_raw_debug)
        st.info(f"✅ Transcript fetched via: **{used_path}**")

        full_text = " ".join([e.get("text", "") for e in entries])
        full_text = clean_transcript_text(full_text)

        if not full_text.strip():
            st.error("❌ Transcript retrieved but empty.")
            st.stop()

        # Stima token grezza (~4 char/token)
        approx_tokens = max(1, len(full_text) // 4)
        st.caption(f"≈ Estimated tokens transcript: ~{approx_tokens:,}")

        # 2) Sintesi
        chapters = []
        with st.spinner("🧠 Generating summary..."):
            if summary_mode == "Standard (auto-chunk)":
                summary = summarize_text(client, full_text, model=model, temperature=temperature, min_words=int(min_words))

            elif summary_mode == "Map-Reduce (token-safe)":
                summary = summarize_text_safe_mapreduce(client, full_text, model, temperature, int(min_words))

            elif summary_mode == "Chapter-based Map-Reduce (if available)":
                chapters = get_chapters(video_id)
                if not chapters:
                    st.warning("Nessun capitolo rilevato. Uso Map-Reduce standard.")
                    summary = summarize_text_safe_mapreduce(client, full_text, model, temperature, int(min_words))
                else:
                    partials = []
                    for ch in chapters:
                        ch_text = text_by_time_window(entries, ch["start_time"], ch["end_time"])
                        if not ch_text.strip():
                            continue
                        part = summarize_text(client, ch_text, model=model, temperature=temperature, min_words=150)
                        partials.append(f"### {ch['title']}\n{part}")
                    final_prompt = (
                        "Merge the chapter summaries into a single cohesive summary (400-600 words), "
                        "keeping logical flow and the most critical insights.\n\n" + "\n\n".join(partials)
                    )
                    final_resp = client.chat.completions.create(
                        model=model,
                        messages=[
                            {"role": "system", "content": "You are a precise editor."},
                            {"role": "user", "content": final_prompt},
                        ],
                        temperature=temperature,
                        max_tokens=1200,
                    )
                    summary = final_resp.choices[0].message.content

            elif summary_mode == "Extractive (no OpenAI)":
                summary = extractive_summary(full_text, max_sentences=20)

        # 3) Download: TXT + DOCX + ZIP
        transcript_bytes = full_text.encode('utf-8')
        transcript_filename = f"transcript_{video_id}.txt"

        # Bottone transcript .txt standalone
        st.download_button(
            "📝 Download transcript (.txt)",
            data=transcript_bytes,
            file_name=transcript_filename,
            mime="text/plain",
            use_container_width=True,
        )

        # DOCX summary + ZIP
        doc = Document()
        doc.add_heading("Riassunto del Video", level=1)
        doc.add_paragraph(summary)
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        summary_filename = f"summary_{video_id}.docx"

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(transcript_filename, transcript_bytes)
            zf.writestr(summary_filename, docx_buffer.getvalue())
        zip_buffer.seek(0)
        zip_filename = f"youtube_summary_{video_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

        st.download_button(
            "📥 Download ZIP (transcript + summary.docx)",
            data=zip_buffer,
            file_name=zip_filename,
            mime="application/zip",
            use_container_width=True,
        )

        with st.expander("👀 Preview summary"):
            st.write(summary)

        if chapters:
            with st.expander("📑 Capitoli rilevati"):
                for ch in chapters:
                    st.write(f"- {ch['title']} ({ch['start_time']}s → {ch['end_time']}s)")

    except TranscriptsDisabled:
        st.error("❌ Transcripts are disabled for this video.")
    except (NoTranscriptFound, NoTranscriptAvailable) as e:
        st.error(f"❌ No transcript found: {e}")
    except Exception as e:
        st.error(f"❌ Error during extraction: {type(e).__name__}: {e}")