import streamlit as st
import pandas as pd
import openai
import re
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ─────────────────────────────────────────────
# PAGE CONFIG
# set_page_config() va chiamato SOLO nell'entry
# point (Home.py). In una app multipage questa
# riga deve restare commentata.
# Decommentare solo per esecuzione standalone.
# ─────────────────────────────────────────────
# st.set_page_config(
#     page_title="LB Content Generator",
#     page_icon="✍️",
#     layout="wide",
# )

# ─────────────────────────────────────────────
# SIDEBAR – solo configurazione globale
# La API key viene salvata in session_state e
# condivisa automaticamente con tutte le pagine.
# ─────────────────────────────────────────────
with st.sidebar:
    st.header("Configurazione Globale")

    st.subheader("🔐 OpenAI API Key")
    _key_input = st.text_input(
        "API Key",
        type="password",
        placeholder="sk-...",
        value=st.session_state.get("openai_api_key", ""),
        help="La chiave viene condivisa tra tutte le pagine dell'app.",
    )
    if _key_input:
        st.session_state["openai_api_key"] = _key_input

    api_key = st.session_state.get("openai_api_key", "")

    if api_key:
        st.success("API Key configurata")
    else:
        st.warning("Inserisci una API Key per usare il tool.")

    st.divider()
    st.caption("NVL Agency – SEO Tools")

# ─────────────────────────────────────────────
# HEADER PAGINA
# ─────────────────────────────────────────────
st.title("✍️ Link Building Content Generator")
st.caption("Genera articoli SEO ottimizzati a partire dal piano di link building")

st.divider()

# ─────────────────────────────────────────────
# COLUMN DETECTION – mapping fuzzy sui nomi colonna
# Permette allo script di funzionare con file Excel
# che usano nomi colonna diversi o varianti linguistiche.
# ─────────────────────────────────────────────

# Colonne "semantiche" da rilevare automaticamente.
# Ogni chiave è il nome interno usato nel resto dello script;
# il valore è una lista di alias accettati (lowercase, senza spazi).
SEMANTIC_COLS = {
    "anchor_text": [
        "anchortext", "anchor", "anchortext", "testo ancoraggio",
        "testoancoraggio", "keyword", "kw", "parola chiave", "parolachiave",
    ],
    "url": [
        "url", "link", "landingpage", "landing page", "landing", "paginadestinazione",
        "pagina destinazione", "destinazione",
    ],
    "tipologia": [
        "tipologiacontenuto", "tipologia contenuto", "tipologia", "tipo",
        "tipocontenuto", "tipo contenuto", "content type", "contenttype",
        "formato", "categoria contenuto",
    ],
    "titolo": [
        "titolo suggerito", "titolosuggerito", "title suggerito",
        "titlesuggerito", "titolo", "title", "h1", "titolo articolo",
        "titoloarticolo", "suggested title",
    ],
    "descrizione": [
        "descrizione/argomento", "descrizione argomento", "descrizioneargomento",
        "descrizione", "argomento", "brief", "topic", "sintesi", "abstract",
        "note", "contenuto", "traccia",
    ],
    # Colonne di contesto aggiuntivo (opzionali, incluse nel prompt se presenti)
    "publisher":  ["publisher", "editore", "sito", "dominio", "website"],
    "numero":     ["#", "num", "numero", "id", "n."],
}

def _normalise(s: str) -> str:
    """Lowercase + rimuove spazi multipli e caratteri speciali per il confronto."""
    import re
    return re.sub(r"[\s\-_/\]+", " ", str(s).lower().strip())

def detect_columns(df_columns: list) -> dict:
    """
    Cerca ogni colonna semantica nel DataFrame.
    Restituisce un dict  {chiave_interna: nome_colonna_reale | None}.
    Strategia:
      1. Match esatto (normalizzato)
      2. Match per contenimento (la colonna contiene l'alias o viceversa)
    """
    norm_cols = {_normalise(c): c for c in df_columns}
    result = {}
    for key, aliases in SEMANTIC_COLS.items():
        found = None
        # Passata 1: match esatto normalizzato
        for alias in aliases:
            norm_alias = _normalise(alias)
            if norm_alias in norm_cols:
                found = norm_cols[norm_alias]
                break
        # Passata 2: match per contenimento
        if not found:
            for alias in aliases:
                norm_alias = _normalise(alias)
                for norm_c, real_c in norm_cols.items():
                    if norm_alias in norm_c or norm_c in norm_alias:
                        found = real_c
                        break
                if found:
                    break
        result[key] = found
    return result

# ─────────────────────────────────────────────
# COLUMN DETECTION – mapping fuzzy sui nomi colonna
# Permette allo script di funzionare con file Excel
# che usano nomi colonna diversi o varianti linguistiche.
# ─────────────────────────────────────────────

# Colonne "semantiche" da rilevare automaticamente.
# Ogni chiave è il nome interno canonico usato nel resto dello script;
# il valore è la lista di alias accettati (verranno normalizzati).
SEMANTIC_COLS = {
    "anchor_text": [
        "anchor text", "anchor", "testo ancoraggio", "testo ancora",
        "anchortext", "keyword", "kw", "parola chiave", "parolachiave",
    ],
    "url": [
        "url", "link", "landing page", "landingpage", "landing",
        "pagina destinazione", "paginadestinazione", "destinazione",
    ],
    "tipologia": [
        "tipologia contenuto", "tipologiacontenuto", "tipologia", "tipo",
        "tipo contenuto", "tipocontenuto", "content type", "contenttype",
        "formato", "categoria contenuto",
    ],
    "titolo": [
        "title suggerito", "titlesuggerito", "titolo suggerito",
        "titolosuggerito", "titolo", "title", "h1",
        "titolo articolo", "titoloarticolo", "suggested title",
    ],
    "descrizione": [
        "descrizione/argomento", "descrizione argomento", "descrizioneargomento",
        "descrizione", "argomento", "brief", "topic", "sintesi",
        "abstract", "note", "contenuto", "traccia",
    ],
    "publisher": ["publisher", "editore", "sito", "dominio", "website"],
    "numero":    ["#", "num", "numero", "id", "n."],
}

REQUIRED_KEYS   = {"anchor_text", "url", "tipologia", "titolo", "descrizione"}
STRUCTURED_KEYS = {"anchor_text", "url", "tipologia", "titolo", "descrizione"}

CANONICAL_LABELS = {
    "numero":      "#",
    "publisher":   "Publisher",
    "anchor_text": "Anchor Text",
    "url":         "URL",
    "tipologia":   "Tipologia Contenuto",
    "titolo":      "Titolo Suggerito",
    "descrizione": "Descrizione / Argomento",
}

CANONICAL_LABELS_PROMPT = {
    "anchor_text": "Anchor Text",
    "url":         "URL di destinazione",
    "tipologia":   "Tipologia Contenuto",
    "titolo":      "Titolo Suggerito",
    "descrizione": "Brief / Argomento",
    "publisher":   "Publisher / Sito",
    "numero":      "#",
}


def _norm(s: str) -> str:
    """Lowercase + collassa spazi/separatori per confronto."""
    return re.sub(r"[\s\-_/\\]+", " ", str(s).lower().strip())


def detect_columns(df_columns: list) -> dict:
    """
    Cerca ogni colonna semantica nella lista di colonne del DataFrame.
    Ritorna {chiave_canonica: nome_colonna_reale | None}.

    Strategia a due passate:
      1. Match esatto (dopo normalizzazione)
      2. Match per contenimento (l'alias è contenuto nel nome colonna o viceversa)
    """
    norm_map = {_norm(c): c for c in df_columns}
    result = {}
    for key, aliases in SEMANTIC_COLS.items():
        found = None
        norm_aliases = [_norm(a) for a in aliases]
        # Passata 1 – match esatto
        for na in norm_aliases:
            if na in norm_map:
                found = norm_map[na]
                break
        # Passata 2 – contenimento
        if not found:
            for na in norm_aliases:
                for nc, rc in norm_map.items():
                    if na in nc or nc in na:
                        found = rc
                        break
                if found:
                    break
        result[key] = found
    return result


# ─────────────────────────────────────────────
# STEP 1 – CARICAMENTO FILE EXCEL
# ─────────────────────────────────────────────
st.subheader("1. Carica il file di pianificazione")

uploaded_file = st.file_uploader(
    "File Excel (.xlsx) — foglio e colonne vengono rilevati automaticamente",
    type=["xlsx"],
    key="main_upload",
)

df_raw  = None
col_map = {}

if uploaded_file:
    try:
        xl = pd.ExcelFile(uploaded_file)
        sheet_names = xl.sheet_names

        # Pre-selezione foglio: preferisce fogli con keyword note nel nome
        preferred = [s for s in sheet_names if any(
            kw in s.lower()
            for kw in ["q1", "q2", "q3", "q4", "pianif", "planning", "link build"]
        )]
        default_sheet = preferred[0] if preferred else sheet_names[0]

        sheet_used = st.selectbox(
            "Foglio da utilizzare",
            options=sheet_names,
            index=sheet_names.index(default_sheet),
            help="Lo script ha pre-selezionato il foglio più probabile.",
        )

        df_raw = pd.read_excel(xl, sheet_name=sheet_used, header=0)
        df_raw.columns = df_raw.columns.str.strip()

        # ── Rilevamento colonne ──────────────────────────────
        col_map = detect_columns(df_raw.columns.tolist())
        titolo_col = col_map.get("titolo")

        if not titolo_col:
            st.error(
                "Nessuna colonna 'titolo' rilevata automaticamente. "
                "Aprire il pannello di mapping qui sotto e selezionarla manualmente."
            )

        # ── UI di verifica / correzione mapping ─────────────
        missing_auto = [CANONICAL_LABELS[k] for k in REQUIRED_KEYS if not col_map.get(k)]
        with st.expander(
            "🔍 Colonne rilevate — verifica o correggi il mapping",
            expanded=bool(missing_auto),
        ):
            st.caption(
                "Lo script ha cercato automaticamente le colonne necessarie. "
                "Le colonne marcate con * sono obbligatorie. "
                "Seleziona manualmente quelle non rilevate o errate."
            )
            all_opts = ["— non presente —"] + df_raw.columns.tolist()
            new_map = {}
            cols_ui = list(CANONICAL_LABELS.items())
            ui_col_a, ui_col_b = st.columns(2)
            for idx, (key, label) in enumerate(cols_ui):
                detected  = col_map.get(key)
                req_mark  = " *" if key in REQUIRED_KEYS else ""
                cur_idx   = all_opts.index(detected) if detected in all_opts else 0
                container = ui_col_a if idx % 2 == 0 else ui_col_b
                with container:
                    chosen = st.selectbox(
                        f"{label}{req_mark}",
                        options=all_opts,
                        index=cur_idx,
                        key=f"colmap_{key}",
                    )
                new_map[key] = None if chosen == "— non presente —" else chosen

            col_map = new_map   # usa il mapping eventualmente corretto dall'operatore

        # ── Validazione ─────────────────────────────────────
        still_missing = [CANONICAL_LABELS[k] for k in REQUIRED_KEYS if not col_map.get(k)]
        if still_missing:
            st.error(
                f"Completa il mapping prima di continuare: "
                f"**{', '.join(still_missing)}**"
            )
            df_raw = None
        else:
            titolo_col = col_map["titolo"]
            df_raw = df_raw[
                df_raw[titolo_col].notna() &
                (df_raw[titolo_col].astype(str).str.strip() != "")
            ].copy().reset_index(drop=True)

            st.success(
                f"Foglio **{sheet_used}** · "
                f"**{len(df_raw)} articoli** trovati · "
                f"colonne mappate correttamente"
            )

    except Exception as e:
        st.error(f"Errore nel caricamento: {e}")
        df_raw = None

# ─────────────────────────────────────────────
# STEP 2 – CONFIGURAZIONE DEL TOOL
# ─────────────────────────────────────────────
st.subheader("2. Configurazione")

with st.expander("⚙️ Modello, prompt e linee guida", expanded=True):

    col_model, col_words = st.columns([2, 1])

    with col_model:
        st.markdown("**🤖 Modello OpenAI**")
        model = st.selectbox(
            "Modello",
            options=[
                # ── GPT-5.x (famiglia attuale) ──────────────────
                "gpt-5.4",
                "gpt-5.4-mini",
                "gpt-5.4-nano",
                "gpt-5.2",
                "gpt-5",
                # ── Reasoning o-series ───────────────────────────
                "o3",
                "o4-mini",
                # ── GPT-4.1 (aprile 2025) ────────────────────────
                "gpt-4.1",
                "gpt-4.1-mini",
                "gpt-4.1-nano",
                # ── Legacy ───────────────────────────────────────
                "gpt-4o",
                "gpt-4o-mini",
                "gpt-4-turbo",
            ],
            index=0,
            label_visibility="collapsed",
            help=(
                "gpt-5.4 / mini / nano: famiglia flagship attuale (2026)\n"
                "gpt-5.2: precedente flagship, ancora disponibile in API\n"
                "gpt-5: alias stabile alla versione 2025-08-07\n"
                "o3 / o4-mini: reasoning avanzato\n"
                "gpt-4.1 / mini / nano: aprile 2025\n"
                "gpt-4o / mini: legacy multimodale"
            ),
        )
        st.caption(
            "**gpt-5.4** — flagship 2026 · "
            "**gpt-5.4-mini/nano** — veloci ed economici · "
            "**o3/o4-mini** — reasoning · "
            "**gpt-4.1** — legacy aprile 2025"
        )

    with col_words:
        st.markdown("**📏 Lunghezza minima**")
        target_words = st.slider(
            "Parole",
            min_value=300,
            max_value=2000,
            value=700,
            step=50,
        )

    st.markdown("---")

    # ── [settore] ─────────────────────────────────────────────────────────
    st.markdown("**🏷️ Settore**")
    st.caption(
        "Valore del segnaposto **[settore]** — uguale per tutti gli articoli del batch. "
        "Sostituisce il campo Tipologia Contenuto come fonte del settore."
    )
    settore_globale = st.text_input(
        "Settore",
        value=st.session_state.get("settore_globale", ""),
        placeholder="Es: skincare e beauty, finanza personale, mobilità urbana…",
        label_visibility="collapsed",
    )
    if settore_globale:
        st.session_state["settore_globale"] = settore_globale

    st.markdown("---")

    # ── Template prompt ───────────────────────────────────────────────────
    st.markdown("**📋 Template Prompt**")
    st.caption(
        "Segnaposto disponibili: **[settore]**, **[argomento]**, **[ancora]**, **[tipologia]** — "
        "sostituiti automaticamente per ogni riga. "
        "**[tipologia]** viene espanso con le istruzioni specifiche configurate qui sotto."
    )
    system_prompt = st.text_area(
        "Template prompt",
        value=(
            "Agisci come un esperto del settore indicato e scrivi un approfondimento "
            "interessante e coinvolgente senza punti elenco ma con sottotitoli e grassetti.\n\n"
            "All'interno del testo deve essere inserita in modo naturale e solo una volta "
            "la frase indicata nei dati, senza modifiche e in grassetto."
        ),
        height=180,
        label_visibility="collapsed",
    )

    st.markdown("---")

    # ── Linee guida / ToV ─────────────────────────────────────────────────
    st.markdown("**📎 Linee Guida / Tone of Voice / Dos & Don'ts**")
    st.caption("Incolla il testo oppure carica un file .txt o .md — verrà aggiunto al prompt di sistema.")

    col_tov, col_upload = st.columns([3, 2])

    with col_tov:
        guidelines_text = st.text_area(
            "Linee guida",
            placeholder=(
                "Es:\n"
                "TONE OF VOICE\n"
                "• Caldo, empatico, mai aggressivo\n"
                "• Usa la seconda persona singolare\n\n"
                "DOS\n"
                "• Frasi brevi e paragrafi aerei\n"
                "• Titoli informativi e keyword-rich\n\n"
                "DON'TS\n"
                "• Superlativi esasperati\n"
                "• Promesse di tipo medico\n"
                "• Menzione di competitor"
            ),
            height=200,
            label_visibility="collapsed",
        )

    with col_upload:
        guidelines_file = st.file_uploader(
            "Carica file linee guida",
            type=["txt", "md"],
            help="Il contenuto del file verrà aggiunto alle linee guida testuali.",
            label_visibility="collapsed",
        )
        if guidelines_file:
            file_content = guidelines_file.read().decode("utf-8", errors="ignore")
            guidelines_text = (guidelines_text or "") + "\n\n" + file_content
            st.info(f"📄 **{guidelines_file.name}** caricato ({len(file_content.split())} parole)")

# ─────────────────────────────────────────────
# STEP 3 – REVISIONE DATI
# ─────────────────────────────────────────────
if df_raw is not None:
    st.subheader("3. Revisione e modifica dei dati")

    # Rinomina le colonne mappate ai nomi canonici
    rename_map = {v: k for k, v in col_map.items() if v}
    df_canonical = df_raw.rename(columns=rename_map)

    # Colonne extra: tutto ciò che non è stato mappato a un canonico
    mapped_real = set(v for v in col_map.values() if v)
    extra_cols  = [c for c in df_raw.columns if c not in mapped_real]
    extra_canonical = [rename_map.get(c, c) for c in extra_cols]

    # Ordine display: canoniche note prima, poi extra
    canonical_order = ["numero", "publisher", "anchor_text", "url", "tipologia", "titolo", "descrizione"]
    canonical_present = [c for c in canonical_order if c in df_canonical.columns]
    display_cols = canonical_present + [c for c in extra_canonical if c not in canonical_present]

    # Column config per data_editor
    col_config = {}
    for c in display_cols:
        label = CANONICAL_LABELS.get(c, c)
        if c == "numero":
            col_config[c] = st.column_config.TextColumn(label, width="small")
        elif c in ("publisher", "tipologia"):
            col_config[c] = st.column_config.TextColumn(label, width="medium")
        elif c in ("url", "anchor_text", "titolo"):
            col_config[c] = st.column_config.TextColumn(label, width="large")
        elif c == "descrizione":
            col_config[c] = st.column_config.TextColumn(label, width="xlarge")
        else:
            col_config[c] = st.column_config.TextColumn(label, width="medium")

    st.info(
        "Modifica direttamente le celle prima di generare. "
        "Le colonne aggiuntive rilevate dal file (in fondo alla tabella) "
        "vengono incluse automaticamente nel prompt di generazione."
    )

    edit_data = df_canonical[display_cols].astype(str).replace("nan", "")
    edited_df = st.data_editor(
        edit_data,
        use_container_width=True,
        num_rows="fixed",
        key="editor",
        column_config=col_config,
    )

    # ─────────────────────────────────────────────
    # STEP 4 – SELEZIONE ARTICOLI
    # ─────────────────────────────────────────────
    st.subheader("4. Seleziona gli articoli da generare")

    all_titles = edited_df["titolo"].tolist()

    col_sel, col_all = st.columns([5, 1])
    with col_all:
        select_all = st.checkbox("Tutti", value=True)

    default_sel = list(range(len(all_titles))) if select_all else []

    selected = st.multiselect(
        "Articoli",
        options=list(range(len(all_titles))),
        default=default_sel,
        format_func=lambda i: f"{i+1}. {all_titles[i][:90]}",
        label_visibility="collapsed",
    )

    if selected:
        st.caption(
            f"**{len(selected)}** articoli selezionati · "
            f"modello: `{model}` · min. {target_words} parole"
        )

    # ─────────────────────────────────────────────
    # HELPER FUNCTIONS
    # ─────────────────────────────────────────────
    def build_user_prompt(row: pd.Series, words: int, template: str,
                           settore: str) -> str:
        """
        Costruisce il prompt per ogni articolo combinando:
          1. Il template statico (uguale per tutti gli articoli del batch)
          2. I dati specifici della riga (titolo, descrizione, anchor, tipologia…)
          3. Le istruzioni strutturali hardcodate (lunghezza, heading, ecc.)
          4. Le istruzioni vincolanti sul link

        Il template non usa segnaposto da sostituire: rimane intatto.
        I dati di riga vengono aggiunti come blocco separato dopo il template.
        """
        tipologia   = row.get("tipologia",    "")
        title       = row.get("titolo",        "")
        descrizione = row.get("descrizione",   "")
        anchor_text = row.get("anchor_text",   "")
        url         = row.get("url",           "")

        # Colonne extra (non strutturate) come contesto aggiuntivo
        extra_lines = []
        for col in row.index:
            if col in STRUCTURED_KEYS or col in ("numero", "publisher"):
                continue
            val = str(row.get(col, "")).strip()
            if val and val.lower() not in ("nan", "none", ""):
                label = CANONICAL_LABELS_PROMPT.get(col, col)
                extra_lines.append(f"  - {label}: {val}")

        # Blocco dati articolo
        dati_articolo = f"DATI DELL'ARTICOLO:\n"
        dati_articolo += f"- Settore: {settore}\n" if settore.strip() else ""
        dati_articolo += f"- Titolo (H1): {title}\n"
        dati_articolo += f"- Argomento / Brief: {descrizione}\n" if descrizione.strip() else ""
        dati_articolo += f"- Ancora da inserire: {anchor_text}\n" if anchor_text.strip() else ""
        # Tipologia: tradotta in istruzione operativa
        TIPOLOGIA_ISTRUZIONI = {
            "contenuto generico":   "Scrivi un articolo informativo e accessibile, adatto a un pubblico generalista.",
            "contenuto verticale":  "Scrivi un articolo specialistico con terminologia tecnica appropriata, rivolto a lettori esperti del settore.",
            "contenuto locale":     "Scrivi un articolo con forte connotazione geografica, con riferimenti locali concreti e informazioni pratiche per chi si trova in quella zona.",
            "listicle":             "Struttura il contenuto come un articolo numerato o a lista tematica, con sezioni ben distinte e titoli evocativi.",
            "how to":               "Scrivi una guida pratica step-by-step, con istruzioni chiare e progressive.",
            "evergreen":            "Scrivi un contenuto senza riferimenti temporali, valido nel tempo, con focus su informazioni stabili e ricercate.",
            "comparativa":          "Struttura il testo come un confronto ragionato tra opzioni, soluzioni o approcci diversi.",
            "approfondimento":      "Scrivi un testo analitico e dettagliato, con focus sulla comprensione profonda del tema.",
        }
        if tipologia.strip():
            tip_key = tipologia.strip().lower()
            tip_istruzione = TIPOLOGIA_ISTRUZIONI.get(tip_key, "")
            if not tip_istruzione:
                # Fallback: match parziale
                for k, v in TIPOLOGIA_ISTRUZIONI.items():
                    if k in tip_key or tip_key in k:
                        tip_istruzione = v
                        break
            if tip_istruzione:
                dati_articolo += f"- Formato richiesto: {tip_istruzione}\n"
            else:
                dati_articolo += f"- Tipologia: {tipologia}\n"
        if extra_lines:
            dati_articolo += "- Altri dati:\n" + "\n".join(extra_lines) + "\n"

        # Prompt finale = template statico + blocco dati articolo
        user_prompt = f"{template.strip()}\n\n{dati_articolo}"

        # Istruzioni strutturali hardcodate (sempre presenti, non configurabili)
        structural = (
            f"\n\nREQUISITI STRUTTURALI:\n"
            f"- Titolo dell'articolo (H1): {title}\n"
            f"- Lunghezza minima: {words} parole\n"
            "- Struttura: introduzione, 3-5 sezioni con sottotitoli (##), conclusione\n"
            "- Usa ## per H2 e ### per H3; niente punti elenco nel corpo\n"
            "- Non aggiungere meta-commenti o note finali sull'articolo stesso"
        )

        # Istruzioni vincolanti sul link (hardcodate)
        link_instruction = ""
        if anchor_text and url and anchor_text.strip() and url.strip():
            _at  = anchor_text
            _url = url
            link_instruction = (
                "\n\nLINK DA INSERIRE:\n"
                f'Il testo \"{_at}\" deve comparire nell\'articolo una sola volta, '
                "in grassetto (**testo**), linkato all'URL indicato.\n\n"
                "Come gestire l'inserimento:\n"
                "Prima di scrivere, individua il concetto del testo a cui l'ancora "
                "e' semanticamente piu' vicina. Costruisci o adatta la frase di quel punto "
                f'in modo che \"{_at}\" vi cada come la parola naturale in quel contesto: '
                "non come elemento aggiunto a fine paragrafo, ma come parte della frase.\n"
                "Se l'ancora e' distante tematicamente, trova il punto di contatto piu' "
                "credibile tra il suo campo semantico e il contenuto, e costruisci la frase "
                "partendo da quel legame.\n\n"
                "Vincoli assoluti:\n"
                "- L'ancora non deve mai essere una frase autonoma o un paragrafo a se.\n"
                "- L'ancora non deve mai essere aggiunta in fondo a un paragrafo gia' concluso.\n"
                "- Non annunciare, introdurre o commentare il link.\n"
                "- Non usare trovare, cliccare, consultare, approfondire, segnalare.\n"
                "- Non usare link, anchor, riferimento, risorsa, fonte vicino al testo linkato.\n"
                "- Non citare l'URL in chiaro nel testo.\n"
                "- Non posizionare nella prima o nell'ultima frase.\n"
                f'Formato: **{_at}**\n'
                f'URL (solo per il link, non scriverlo nel testo): {_url}'
            )

        return user_prompt + structural + link_instruction

    def build_system_prompt(guidelines: str) -> str:
        """Prompt di sistema: solo linee guida/ToV. Il template è nel user prompt."""
        if guidelines and guidelines.strip():
            return (
                "Sei un copywriter professionista. "
                "Rispetta scrupolosamente le seguenti linee guida in tutto ciò che scrivi:\n\n"
                + guidelines.strip()
            )
        return "Sei un copywriter professionista. Scrivi in italiano."

    def call_openai(client, sys_prompt: str, user_prompt: str, mdl: str) -> str:
        # Modelli che richiedono max_completion_tokens invece di max_tokens:
        #   - o-series  (o1, o3, o4-mini, ...)
        #   - GPT-5.x   (gpt-5, gpt-5.1, gpt-5.2, gpt-5.4, ...)
        # I modelli "o" non supportano inoltre il ruolo system separato.
        use_completion_tokens = mdl.startswith("o") or mdl.startswith("gpt-5")

        if mdl.startswith("o"):
            messages = [{"role": "user", "content": f"{sys_prompt}\n\n---\n\n{user_prompt}"}]
        else:
            messages = [
                {"role": "system", "content": sys_prompt},
                {"role": "user",   "content": user_prompt},
            ]

        kwargs = {"model": mdl, "messages": messages}
        if use_completion_tokens:
            kwargs["max_completion_tokens"] = 3000
        else:
            kwargs["max_tokens"] = 3000

        response = client.chat.completions.create(**kwargs)
        return response.choices[0].message.content

    def _add_hyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        wr = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
        clr = OxmlElement("w:color"); clr.set(qn("w:val"), "1A73E8"); rPr.append(clr)
        rPr.append(OxmlElement("w:b"))
        wr.append(rPr)
        t = OxmlElement("w:t"); t.text = text; wr.append(t)
        hyperlink.append(wr)
        paragraph._p.append(hyperlink)

    def markdown_to_docx(doc: Document, text: str, url: str, anchor_text: str = ""):
        bold_pattern = re.compile(r'\*\*(.+?)\*\*')
        anchor_norm  = anchor_text.strip().lower()

        for line in text.split("\n"):
            line = line.rstrip()
            if not line:
                doc.add_paragraph("")
                continue
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            else:
                p    = doc.add_paragraph()
                last = 0
                for m in bold_pattern.finditer(line):
                    if line[last:m.start()]:
                        p.add_run(line[last:m.start()])
                    bold_text = m.group(1)
                    if anchor_norm and bold_text.strip().lower() == anchor_norm and url:
                        _add_hyperlink(p, bold_text, url)
                    else:
                        p.add_run(bold_text).bold = True
                    last = m.end()
                if line[last:]:
                    p.add_run(line[last:])


    def _set_cell(cell, text, bold=False, is_url=False, bg_hex=None, font_size=9):
        cell.text = ""
        p   = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(text)
        run.bold      = bold
        run.font.name = "Arial"
        run.font.size = Pt(font_size)
        if is_url:
            run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
            run.underline = True
        if bg_hex:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  bg_hex)
            tcPr.append(shd)

    def create_word_doc(row: pd.Series, generated_text: str) -> BytesIO:
        doc = Document()

        # Margini
        for section in doc.sections:
            section.top_margin    = Inches(0.9)
            section.bottom_margin = Inches(0.9)
            section.left_margin   = Inches(1.1)
            section.right_margin  = Inches(1.1)

        doc.styles["Normal"].font.name = "Arial"
        doc.styles["Normal"].font.size = Pt(11)

        publisher = str(row.get("publisher",   "") or "")
        tipologia = str(row.get("tipologia",   "") or "")
        anchor    = str(row.get("anchor_text", "") or "")
        url       = str(row.get("url",         "") or "")
        titolo    = str(row.get("titolo",       "") or "")

        # ── Tabella di riepilogo ──────────────────────────────────
        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"

        meta = [
            ("Publisher",   publisher, False),
            ("Tipologia",   tipologia, False),
            ("Anchor Text", anchor,    False),
            ("URL",         url,       True),
        ]

        HDR_BG = "1F3864"
        VAL_BG = "F5F7FA"

        for i, (label, value, is_url) in enumerate(meta):
            lc = table.cell(i, 0)
            vc = table.cell(i, 1)
            _set_cell(lc, label, bold=True, bg_hex=HDR_BG)
            lc.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell(vc, value, is_url=is_url, bg_hex=VAL_BG)

        # Larghezze colonne
        for row_t in table.rows:
            row_t.cells[0].width = Inches(3.5 / 2.54)
            row_t.cells[1].width = Inches(13.0 / 2.54)

        # ── Separatore ───────────────────────────────────────────
        doc.add_paragraph("")

        # ── H1 titolo ────────────────────────────────────────────
        h1 = doc.add_heading(titolo, level=1)
        h1.paragraph_format.space_before = Pt(6)
        h1.paragraph_format.space_after  = Pt(12)

        # ── Corpo ────────────────────────────────────────────────
        markdown_to_docx(doc, generated_text, url, anchor)

        # ── Footer ───────────────────────────────────────────────
        doc.add_paragraph("")
        fp = doc.add_paragraph()
        fp.paragraph_format.space_before = Pt(10)
        fr = fp.add_run(
            f"Generato con LB Content Generator \u2013 NVL Agency \u00b7 modello: {model}"
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ─────────────────────────────────────────────
    # STEP 5 – GENERA
    # ─────────────────────────────────────────────
    st.subheader("5. Genera i contenuti")

    # ── Preview prompt ────────────────────────────────────────────────
    if selected and df_raw is not None:
        with st.expander("🔍 Anteprima prompt — riga selezionata", expanded=False):
            preview_idx = st.selectbox(
                "Articolo da visualizzare",
                options=selected,
                format_func=lambda i: f"{i+1}. {edited_df.iloc[i].get('titolo', '')[:70]}",
                key="preview_select",
            )
            preview_row    = edited_df.iloc[preview_idx]
            preview_sys    = build_system_prompt(guidelines_text)
            preview_user   = build_user_prompt(preview_row, target_words, system_prompt, settore_globale)
            col_ps, col_pu = st.columns(2)
            with col_ps:
                st.markdown("**System prompt**")
                st.text_area("sys", preview_sys, height=200, label_visibility="collapsed", key="prev_sys", disabled=True)
            with col_pu:
                st.markdown("**User prompt**")
                st.text_area("usr", preview_user, height=200, label_visibility="collapsed", key="prev_usr", disabled=True)
            total_chars = len(preview_sys) + len(preview_user)
            est_tokens  = total_chars // 4
            st.caption(f"Stima token prompt: ~{est_tokens:,} · {total_chars:,} caratteri")

    # ── Stima costo batch ─────────────────────────────────────────────
    COST_PER_1K = {
        "gpt-5.4":       (0.005, 0.020),
        "gpt-5.4-mini":  (0.001, 0.004),
        "gpt-5.4-nano":  (0.0005, 0.002),
        "gpt-5.2":       (0.004, 0.016),
        "gpt-5":         (0.004, 0.016),
        "o3":            (0.010, 0.040),
        "o4-mini":       (0.002, 0.008),
        "gpt-4.1":       (0.002, 0.008),
        "gpt-4.1-mini":  (0.0004, 0.0016),
        "gpt-4.1-nano":  (0.0001, 0.0004),
        "gpt-4o":        (0.0025, 0.010),
        "gpt-4o-mini":   (0.00015, 0.0006),
        "gpt-4-turbo":   (0.010, 0.030),
    }
    if selected and model in COST_PER_1K:
        in_cost, out_cost = COST_PER_1K[model]
        # Stima: ~800 token in, ~1000 token out per articolo
        est_in  = 800  * len(selected) / 1000
        est_out = 1000 * len(selected) / 1000
        est_total = est_in * in_cost + est_out * out_cost
        st.caption(
            f"💰 Costo stimato batch: **${est_total:.3f}** "
            f"({len(selected)} articoli × modello `{model}`) — stima indicativa"
        )

    col_btn, col_info = st.columns([2, 5])
    with col_btn:
        generate_btn = st.button(
            "🚀 Genera contenuti",
            type="primary",
            disabled=(not api_key or not selected),
            use_container_width=True,
        )
    with col_info:
        if not api_key:
            st.warning("Configura la API Key nella sidebar per procedere.")
        elif not selected:
            st.warning("Seleziona almeno un articolo nello step 4.")

    if generate_btn:
        client     = openai.OpenAI(api_key=api_key)
        sys_prompt = build_system_prompt(guidelines_text)

        generated_docs = {}
        errors         = []

        progress = st.progress(0, text="Avvio generazione...")
        status   = st.empty()

        for i, row_idx in enumerate(selected):
            row   = edited_df.iloc[row_idx]
            title = row.get("titolo", f"Articolo {row_idx+1}")
            progress.progress(
                int(i / len(selected) * 100),
                text=f"({i+1}/{len(selected)}) {title[:55]}...",
            )
            status.info(f"⏳ Generando: **{title}**")

            try:
                user_p = build_user_prompt(row, target_words, system_prompt, settore_globale)
                text   = call_openai(client, sys_prompt, user_p, model)
                # Controllo lunghezza — se sotto 80% della soglia, chiede espansione
                word_count = len(text.split())
                if word_count < target_words * 0.80:
                    status.warning(f"⚠️ {title[:40]}… ({word_count} parole, sotto soglia) — espansione in corso…")
                    expand_p = (
                        f"Il testo seguente è troppo breve ({word_count} parole su {target_words} richieste). "
                        f"Espandilo fino ad almeno {target_words} parole aggiungendo dettagli, "
                        "esempi e approfondimenti coerenti con il contenuto esistente. "
                        "Non alterare il titolo, la struttura dei paragrafi o il link già inserito.\n\n"
                        + text
                    )
                    text = call_openai(client, sys_prompt, expand_p, model)
                doc_buf = create_word_doc(row, text)
                safe    = re.sub(r"[^\w\-_]", "_", str(title))[:60]
                wc_final = len(text.split())
                fname_key = f"{row_idx+1:02d}_{safe}.docx"
                generated_docs[fname_key] = (doc_buf, wc_final)
            except Exception as e:
                errors.append(f"Riga {row_idx+1} \u2013 {title}: {e}")

        progress.progress(100, text="Completato!")
        status.empty()

        if errors:
            st.error("Errori durante la generazione:\n" + "\n".join(errors))

        if generated_docs:
            wc_info = {k: wc for k, (_, wc) in generated_docs.items()}
            low_wc  = {k: wc for k, wc in wc_info.items() if wc < target_words}
            st.success(f"✅ {len(generated_docs)} articoli generati con successo!")
            if low_wc:
                st.warning("Articoli ancora sotto soglia dopo espansione: " +
                    ", ".join(f"{k} ({v} parole)" for k, v in low_wc.items()))

            st.subheader("6. Download")

            if len(generated_docs) == 1:
                fname, (buf, _wc) = next(iter(generated_docs.items()))
                st.caption(f"📝 {wc_info.get(fname, '?')} parole")
                st.download_button(
                    label=f"📄 Scarica {fname}",
                    data=buf,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            else:
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
                    for fname, (buf, _wc) in generated_docs.items():
                        zf.writestr(fname, buf.read())
                zip_buf.seek(0)

                st.download_button(
                    label=f"📦 Scarica tutti ({len(generated_docs)} articoli) — .zip",
                    data=zip_buf,
                    file_name="contenuti_generati.zip",
                    mime="application/zip",
                    type="primary",
                )

                st.write("Oppure scarica i file singolarmente:")
                dl_cols = st.columns(min(len(generated_docs), 3))
                for j, (fname, (buf, _wc)) in enumerate(generated_docs.items()):
                    buf.seek(0)
                    with dl_cols[j % 3]:
                        st.caption(f"📝 {wc_info.get(fname, '?')} parole")
                        st.download_button(
                            label=f"📄 {fname[:40]}",
                            data=buf,
                            file_name=fname,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_{fname}",
                        )
