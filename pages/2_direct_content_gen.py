import streamlit as st
from io import BytesIO
from docx import Document
from openai import OpenAI, BadRequestError

# =========================
# 📌 Configurazione pagina
# =========================
st.set_page_config(page_title="AI Content Generator", layout="wide")
st.title("📝 AI Content Generator - Articles & Recipes")

# =========================
# ⚙️ Utility & Token Budget
# =========================
def approx_tokens(text: str) -> int:
    """Stima grezza dei token: ~4 caratteri per token."""
    if not text:
        return 0
    return max(1, len(text) // 4)

# Solo modelli con 128k contesto
CONTEXT_LIMITS = {
    "gpt-4o-mini": 128000,
    "gpt-4o": 128000,
    "gpt-5": 128000,
}

def compute_effective_max_tokens(
    model: str,
    system_prompt: str,
    user_prompt: str,
    desired_max_tokens: int
) -> int:
    """Calcola un max_tokens che non superi la context window del modello."""
    ctx_limit = CONTEXT_LIMITS.get(model, 128000)
    sys_t = approx_tokens(system_prompt)
    usr_t = approx_tokens(user_prompt)
    overhead = 100  # margine per ruoli/metadata
    max_allowed_completion = max(256, ctx_limit - (sys_t + usr_t + overhead))
    return max(256, min(desired_max_tokens, max_allowed_completion))

def safe_chat_completion(
    client: OpenAI,
    system_prompt: str,
    user_prompt: str,
    model: str = "gpt-4o-mini",
    temperature: float = 0.7,
    desired_max_tokens: int = 4000
) -> str:
    """
    Invio robusto a Chat Completions con:
    - calcolo dinamico max_tokens;
    - retry automatico se finish_reason == 'length';
    - fallback progressivi senza rompere l'app.
    """
    eff_max_tokens = compute_effective_max_tokens(model, system_prompt, user_prompt, desired_max_tokens)

    attempt_budgets = [
        eff_max_tokens,
        min(eff_max_tokens * 2, 8000),
        min(eff_max_tokens * 3, 12000),
    ]

    last_err = None

    for mt in attempt_budgets:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=temperature,
                max_tokens=mt,
            )

            choice = resp.choices[0]
            content = (choice.message.content or "").strip()
            finish_reason = (choice.finish_reason or "stop").lower()

            # Se il modello ha troncato per lunghezza, prova con più token nel prossimo ciclo
            if finish_reason == "length":
                continue

            # Heuristica extra: se finisce male e abbiamo ancora margine, riproviamo
            if content and content[-1] in {",", ";", ":"} and mt < 12000:
                continue

            return content

        except BadRequestError as e:
            last_err = e
            # prompt troppo lungo o simili → si passa al budget successivo (più basso)
            continue
        except Exception as e:
            last_err = e
            continue

    # Ultimo tentativo esplicito per completare integralmente
    try:
        resp2 = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": (
                        user_prompt
                        + "\n\nCompleta il testo in modo integrale, senza riassumere e senza troncare."
                    ),
                },
            ],
            temperature=temperature,
            max_tokens=min(12000, eff_max_tokens * 3),
        )
        return (resp2.choices[0].message.content or "").strip()
    except Exception as e:
        raise RuntimeError(
            "Il prompt complessivo è troppo lungo per la finestra del modello anche dopo i retry. "
            "Riduci la 'Fonte autorevole', il numero minimo di parole o il numero di paragrafi."
        ) from (last_err or e)

# =========================
# 🔐 API Key + Modello (client riusato)
# =========================
def get_openai_client_and_model():
    with st.sidebar:
        st.subheader("🔐 OpenAI")
        api_key = st.text_input("Insert your OpenAI API KEY", type="password")

        model = st.selectbox(
            "Seleziona il modello",
            options=["gpt-4o-mini", "gpt-4o", "gpt-5"],
            index=0,
            help=(
                "Scegli il modello:\n"
                "- gpt-4o-mini → più economico, ottimo per testi brevi/medi\n"
                "- gpt-4o → bilanciato, buona qualità\n"
                "- gpt-5 → massima qualità, ideale per testi lunghi e complessi"
            ),
        )

        if model == "gpt-5":
            st.info("🚀 Stai usando GPT-5: massima qualità e contesto da 128k token.")

        if not api_key:
            st.warning("⚠️ Inserisci la tua API KEY per continuare.")
            st.stop()

    # ✅ Riuso del client per evitare troppe connessioni / file aperti
    if "openai_client" not in st.session_state or st.session_state.get("openai_api_key") != api_key:
        st.session_state["openai_client"] = OpenAI(api_key=api_key)
        st.session_state["openai_api_key"] = api_key

    client = st.session_state["openai_client"]
    return client, model

client, selected_model = get_openai_client_and_model()

# =========================
# 🧭 Tabs
# =========================
tab1, tab2 = st.tabs(["📝 Articles", "🍽 Recipes"])

# =========================
# 📝 TAB ARTICOLI
# =========================
with tab1:
    st.header("📄 Generate Articles")

    with st.form("article_form"):
        title = st.text_input("Title of the article")
        seo_title = st.text_input("SEO Title")
        meta_description = st.text_area("Meta Description")
        min_words = st.number_input(
            "Minimum number of words:",
            min_value=100,
            value=800,
            step=100
        )
        tone_of_voice = st.selectbox(
            "Tone of Voice",
            ["Informale", "Professionale", "Formale"]
        )
        source_text = st.text_area(
            "Fonte autorevole (non verrà citata direttamente)"
        )

        st.subheader("📝 Paragraphs (max 5)")
        paragraphs = []
        for i in range(1, 6):
            col1, col2 = st.columns([1, 2])
            with col1:
                p_title = st.text_input(
                    f"Paragrafo {i} - Titolo",
                    key=f"title_{i}"
                )
            with col2:
                p_desc = st.text_area(
                    f"Paragrafo {i} - Descrizione",
                    key=f"desc_{i}"
                )

            if p_title and p_desc:
                paragraphs.append((p_title, p_desc))

        submit_article = st.form_submit_button("Generate Article")

    if submit_article:
        if not title or not seo_title or not meta_description or not paragraphs:
            st.error("❌ Please fill out all required fields.")
        else:
            st.success(f"📝 Generating article: {title}...")

            min_words_per_para = max(120, int(min_words // len(paragraphs)))
            generated_sections = []

            for p_title, p_desc in paragraphs:
                prompt = f"""
Scrivi un paragrafo con il titolo "{p_title}".
Il tono deve essere {tone_of_voice.lower()}.
Il contenuto deve sviluppare questa idea: "{p_desc}".
Deve contenere almeno {min_words_per_para} parole.
La fonte seguente può servire come contesto (non citarla, non copiarla, non parafrasarla direttamente):
{source_text}
                """.strip()

                target_tokens = int(min_words_per_para * 1.6)

                try:
                    content = safe_chat_completion(
                        client=client,
                        system_prompt="Sei un esperto redattore SEO.",
                        user_prompt=prompt,
                        model=selected_model,
                        temperature=0.7,
                        desired_max_tokens=min(6000, max(800, target_tokens)),
                    )
                except Exception as e:
                    st.error(f"❌ Errore durante la generazione del paragrafo '{p_title}': {e}")
                    content = ""

                generated_sections.append((p_title, content))

            # 📄 Anteprima
            st.subheader("📖 Anteprima Articolo")
            st.markdown(f"# {title}")
            st.markdown(
                f"**SEO Title**: {seo_title}\n\n**Meta Description**: {meta_description}"
            )

            for p_title, p_text in generated_sections:
                if p_text:
                    st.markdown(f"## {p_title}\n\n{p_text}")
                else:
                    st.markdown(f"## {p_title}\n\n*Paragrafo non generato per errore.*")

            # 📁 Export Word (BytesIO)
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"SEO Title: {seo_title}")
            doc.add_paragraph(f"Meta Description: {meta_description}")
            doc.add_paragraph(f"Tone of Voice: {tone_of_voice}")

            for p_title, p_text in generated_sections:
                doc.add_heading(p_title, level=2)
                doc.add_paragraph(p_text if p_text else "[Non disponibile]")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download Article (.docx)",
                data=buffer,
                file_name=f"{title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# =========================
# 🍽 TAB RICETTE
# =========================
with tab2:
    st.header("🍽 Generate Recipes")

    with st.form("recipe_form"):
        recipe_title = st.text_input("Recipe Title")
        min_words_recipe = st.number_input(
            "Minimum number of words:",
            min_value=100,
            value=500,
            step=50
        )
        recipe_tone = st.selectbox(
            "Tone of Voice",
            ["Informale", "Professionale", "Formale"],
            key="recipe_tone",
        )
        recipe_source = st.text_area(
            "Fonte autorevole (non verrà citata direttamente)",
            key="recipe_source",
        )
        ingredients = st.text_area(
            "Insert ingredients (one per line)",
            key="ingredients",
        )
        preparation_desc = st.text_area(
            "Describe the preparation process",
            key="preparation_desc",
        )

        submit_recipe = st.form_submit_button("Generate Recipe")

    if submit_recipe:
        if not recipe_title or not ingredients or not preparation_desc:
            st.error("❌ Please fill out all required fields.")
        else:
            st.success(f"🍽 Generating recipe: {recipe_title}...")

            intro_prompt = f"""
Scrivi una breve introduzione per la ricetta "{recipe_title}".
Il tono deve essere {recipe_tone.lower()}.
La fonte è solo per contesto (non citarla, non copiarla, non parafrasarla direttamente):
{recipe_source}
            """.strip()

            prep_prompt = f"""
Scrivi un paragrafo dettagliato per spiegare come preparare "{recipe_title}".
Deve contenere almeno {min_words_recipe} parole.
Gli ingredienti sono:
{ingredients}

Il tono deve essere {recipe_tone.lower()}.
{preparation_desc}
            """.strip()

            target_tokens_recipe = int(min_words_recipe * 1.6)

            # Introduzione
            try:
                recipe_intro = safe_chat_completion(
                    client=client,
                    system_prompt="Sei un esperto di cucina appassionato di temi quali sostenibilità, territorialità e stagionalità degli ingredienti, degli alimenti e dei processi di preparazione.",
                    user_prompt=intro_prompt,
                    model=selected_model,
                    temperature=0.7,
                    desired_max_tokens=1200,
                )
            except Exception as e:
                st.error(f"❌ Errore durante la generazione dell'introduzione: {e}")
                recipe_intro = ""

            # Preparazione
            try:
                recipe_prep = safe_chat_completion(
                    client=client,
                    system_prompt="Sei un esperto di cucina appassionato di temi quali sostenibilità, territorialità e stagionalità degli ingredienti, degli alimenti e dei processi di preparazione.",
                    user_prompt=prep_prompt,
                    model=selected_model,
                    temperature=0.7,
                    desired_max_tokens=min(8000, max(800, target_tokens_recipe)),
                )
            except Exception as e:
                st.error(f"❌ Errore durante la generazione della preparazione: {e}")
                recipe_prep = ""

            # 📄 Anteprima
            st.subheader("📖 Anteprima Ricetta")
            st.markdown(f"# {recipe_title}")
            if recipe_intro:
                st.markdown(recipe_intro)

            st.markdown("## Ingredienti")
            st.markdown(
                "\n".join(
                    f"- {i.strip()}"
                    for i in ingredients.strip().split("\n")
                    if i.strip()
                )
            )

            st.markdown("## Preparazione")
            st.markdown(
                recipe_prep
                if recipe_prep
                else "*Preparazione non disponibile per errore.*"
            )

            # 📁 Export Word (BytesIO)
            doc = Document()
            doc.add_heading(recipe_title, level=1)
            doc.add_paragraph(f"Tone of Voice: {recipe_tone}")
            if recipe_intro:
                doc.add_paragraph(recipe_intro)

            doc.add_heading("Ingredienti", level=2)
            for ing in ingredients.strip().split("\n"):
                ing = ing.strip()
                if ing:
                    doc.add_paragraph(f"• {ing}")

            doc.add_heading("Preparazione", level=2)
            doc.add_paragraph(
                recipe_prep if recipe_prep else "[Non disponibile]"
            )

            buffer_recipe = BytesIO()
            doc.save(buffer_recipe)
            buffer_recipe.seek(0)

            st.download_button(
                "📥 Download Recipe (.docx)",
                data=buffer_recipe,
                file_name=f"{recipe_title.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )